from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "apps/web/src/pages/inventory/inventory-manual-content.json"
PUBLIC = ROOT / "apps/web/public/panduan/inventory-sales"
OUTPUT = PUBLIC / "Panduan-Pengguna-eBisnis-Inventory-Sales.docx"
IMAGES = PUBLIC / "images"

INK = RGBColor(15, 23, 42)
TEAL = RGBColor(15, 118, 110)
MUTED = RGBColor(71, 85, 105)
PALE = "ECFDF5"
LINE = "CBD5E1"


def shade(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    element = OxmlElement("w:shd")
    element.set(qn("w:fill"), fill)
    properties.append(element)


def set_cell_text(cell, value: str, bold: bool = False, color: RGBColor = INK) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(value)
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(8.5)
    run.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, end])


def configure(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color in (
        ("Title", 31, INK),
        ("Heading 1", 20, INK),
        ("Heading 2", 14, TEAL),
        ("Heading 3", 11, INK),
    ):
        style = styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)

    if "Caption Custom" not in styles:
        caption = styles.add_style("Caption Custom", WD_STYLE_TYPE.PARAGRAPH)
        caption.font.name = "Aptos"
        caption.font.size = Pt(8)
        caption.font.italic = True
        caption.font.color.rgb = MUTED
        caption.paragraph_format.space_after = Pt(8)

    header = section.header.paragraphs[0]
    header.text = "eBisnis Inventory / Sales  |  Panduan Pengguna"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(8)
        run.font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Caruban Medika Nusantara  •  Versi 1.0  •  Halaman ")
    run.font.name = "Aptos"
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED
    add_page_field(footer)


def add_cover(document: Document, data: dict) -> None:
    for _ in range(3):
        document.add_paragraph()
    eyebrow = document.add_paragraph()
    eyebrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = eyebrow.add_run("PANDUAN RESMI • WEB • WINDOWS • ANDROID FLUTTER")
    run.bold = True
    run.font.name = "Aptos"
    run.font.size = Pt(9)
    run.font.color.rgb = TEAL

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(data["meta"]["title"])
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(16)
    run = subtitle.add_run(data["meta"]["subtitle"])
    run.font.name = "Aptos"
    run.font.size = Pt(15)
    run.font.color.rgb = MUTED

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(IMAGES / "alur-end-to-end.png"), width=Inches(6.1))

    meta = document.add_table(rows=2, cols=3)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.autofit = False
    values = [
        ("VERSI", data["meta"]["version"]),
        ("DIPERBARUI", data["meta"]["updated"]),
        ("PENERBIT", data["meta"]["owner"]),
    ]
    for index, (label, value) in enumerate(values):
        meta.columns[index].width = Inches(2.0)
        set_cell_text(meta.cell(0, index), label, True, TEAL)
        set_cell_text(meta.cell(1, index), value, True)
        shade(meta.cell(0, index), PALE)
    document.add_paragraph()
    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("Untuk pemilik, admin, sales, gudang, dan pelanggan terdaftar")
    run.font.name = "Aptos"
    run.font.size = Pt(10)
    run.font.color.rgb = MUTED
    document.add_page_break()


def add_bullet(document: Document, text: str, numbered: bool = False) -> None:
    paragraph = document.add_paragraph(style="List Number" if numbered else "List Bullet")
    paragraph.paragraph_format.left_indent = Cm(0.5)
    paragraph.paragraph_format.first_line_indent = Cm(-0.25)
    paragraph.add_run(text)


def add_callout(document: Document, label: str, text: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, "FFF7ED")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(f"{label}: ")
    run.bold = True
    run.font.color.rgb = RGBColor(154, 52, 18)
    paragraph.add_run(text)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_front_matter(document: Document, data: dict) -> None:
    document.add_heading("Cara Menggunakan Panduan Ini", level=1)
    document.add_paragraph(
        "Panduan ini memakai pendekatan berbasis pekerjaan. Mulailah dari peran Anda, ikuti alur transaksi, lalu gunakan bab laporan dan kontrol saat diperlukan. Tampilan dapat berbeda sedikit antarperangkat, tetapi data dan aturan bisnisnya sama."
    )
    document.add_heading("Daftar Isi", level=2)
    for chapter in data["chapters"]:
        add_bullet(document, chapter["title"])
    add_bullet(document, "Matriks paritas 48 layar aplikasi legacy")
    add_bullet(document, "Glosarium dan checklist operasional")

    document.add_heading("Prinsip Sistem", level=2)
    for item in data["principles"]:
        add_bullet(document, item)

    document.add_heading("Peran dan Batas Akses", level=2)
    table = document.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for index, header in enumerate(("Peran", "Tujuan", "Platform", "Batas")):
        set_cell_text(table.rows[0].cells[index], header, True, RGBColor(255, 255, 255))
        shade(table.rows[0].cells[index], "0F766E")
    for role in data["roles"]:
        cells = table.add_row().cells
        for index, value in enumerate((role["role"], role["purpose"], role["platform"], role["limits"])):
            set_cell_text(cells[index], value, index == 0)

    document.add_heading("Mulai Cepat", level=2)
    for item in data["quickStart"]:
        add_bullet(document, item, numbered=True)
    document.add_page_break()


def add_image(document: Document, filename: str, alt: str) -> None:
    path = IMAGES / filename
    if not path.exists():
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Inches(6.0))
    caption = document.add_paragraph(style="Caption Custom")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.add_run(alt)


def add_chapters(document: Document, data: dict) -> None:
    illustration_map = {
        "mengenal-sistem": ("alur-end-to-end.png", "Alur terintegrasi dari pelanggan dan sales hingga gudang, admin, dan pemilik."),
        "android-sales": ("sales-android.png", "Sales menggunakan Android Flutter untuk katalog, stok, customer, dan order lapangan."),
        "navigasi-web": ("kontrol-web-desktop.png", "Workspace Web dan Desktop untuk kontrol operasional serta laporan pemilik."),
    }
    for chapter in data["chapters"]:
        document.add_heading(chapter["title"], level=1)
        intro = document.add_paragraph()
        run = intro.add_run(chapter["summary"])
        run.bold = True
        run.font.color.rgb = MUTED
        if chapter["id"] in illustration_map:
            add_image(document, *illustration_map[chapter["id"]])
        for section in chapter["sections"]:
            document.add_heading(section["title"], level=2)
            for text in section.get("paragraphs", []):
                document.add_paragraph(text)
            for text in section.get("bullets", []):
                add_bullet(document, text)
            for text in section.get("steps", []):
                add_bullet(document, text, numbered=True)
            if section.get("warning"):
                add_callout(document, "Perhatian", section["warning"])
        document.add_paragraph()


def add_parity_and_glossary(document: Document, data: dict) -> None:
    document.add_page_break()
    document.add_heading("Matriks Paritas 48 Layar Legacy", level=1)
    document.add_paragraph(
        "Fungsi aplikasi DBF lama dipertahankan sebagai kemampuan bisnis dan dipadukan dengan kontrol ERP modern. Matriks ini menjadi checklist penerimaan pengguna, bukan instruksi untuk meniru antarmuka lama."
    )
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, header in enumerate(("Kelompok", "Jumlah", "Fitur yang dipetakan")):
        set_cell_text(table.rows[0].cells[index], header, True, RGBColor(255, 255, 255))
        shade(table.rows[0].cells[index], "0F766E")
    for group in data["parityGroups"]:
        cells = table.add_row().cells
        set_cell_text(cells[0], group["group"], True)
        set_cell_text(cells[1], str(group["count"]))
        set_cell_text(cells[2], "; ".join(group["features"]))
    for row in table.rows:
        for cell in row.cells:
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(7.5)

    document.add_heading("Akses Panduan Publik", level=1)
    document.add_paragraph("Baca online: https://inventory.ebisnis.id/panduan/inventory-sales")
    document.add_paragraph("Unduhan Word dan PDF tersedia pada halaman yang sama. Panduan ringkas offline juga tersedia melalui tab Panduan pada aplikasi Android Flutter.")
    add_callout(document, "Catatan keamanan", "Jangan pernah mengirim kata sandi, token, data pasien, atau data rahasia customer ketika meminta bantuan.")

    document.add_page_break()
    document.add_heading("Glosarium", level=1)
    glossary = document.add_table(rows=1, cols=2)
    glossary.style = "Table Grid"
    glossary.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, header in enumerate(("Istilah", "Arti operasional")):
        set_cell_text(glossary.rows[0].cells[index], header, True, RGBColor(255, 255, 255))
        shade(glossary.rows[0].cells[index], "0F766E")
    for item in data["glossary"]:
        cells = glossary.add_row().cells
        set_cell_text(cells[0], item["term"], True)
        set_cell_text(cells[1], item["definition"])
    for row in glossary.rows:
        for cell in row.cells:
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(7.5)


def main() -> None:
    PUBLIC.mkdir(parents=True, exist_ok=True)
    data = json.loads(SOURCE.read_text(encoding="utf-8-sig"))
    document = Document()
    configure(document)
    add_cover(document, data)
    add_front_matter(document, data)
    add_chapters(document, data)
    add_parity_and_glossary(document, data)
    properties = document.core_properties
    properties.title = data["meta"]["title"]
    properties.subject = data["meta"]["scope"]
    properties.author = data["meta"]["owner"]
    properties.keywords = "inventory, sales, flutter, web, android, ERP, user manual"
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()

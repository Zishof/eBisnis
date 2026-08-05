from __future__ import annotations

import csv
import json
import os
import re
import tempfile
from pathlib import Path
from textwrap import shorten
from zipfile import ZipFile

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
PUBLIC = ROOT / "apps/web/public/panduan/inventory-sales"
MODERN_IMAGES = PUBLIC / "images/transisi/baru"
DOCX = PUBLIC / "Panduan-Transisi-48-Layar-eBisnis-Inventory-Sales.docx"
INDEX = ROOT / "docs/user-manual/inventory-transition-index.json"
REPORT = ROOT / "docs/user-manual/inventory-transition-word-count.json"
PARITY_SOURCE = ROOT / "apps/api/src/modules/tenant/sales-inventory-parity.catalog.ts"

PACKAGE = Path(os.environ.get("INVENTORY_LEGACY_PACKAGE", Path.home() / "Downloads/Paket_Dokumentasi_Sales_Inventory"))
LEGACY_DOCX = Path(os.environ.get("INVENTORY_LEGACY_MANUAL_DOCX", PACKAGE / "User_Manual_Sales_Inventory_Komprehensif.docx"))
MATRIX = Path(os.environ.get("INVENTORY_LEGACY_MATRIX", PACKAGE / "Matriks_Paritas_48_Layar.csv"))

INK = RGBColor(15, 23, 42)
TEAL = RGBColor(15, 118, 110)
MUTED = RGBColor(71, 85, 105)
LIGHT = RGBColor(241, 245, 249)


DOMAIN_CONFIG = {
    "Master Supplier": ("Supplier", ["Kode", "Nama supplier", "Wilayah", "Termin", "Status"], [
        ["SUP-001", "PT Anugerah Sehat", "Cirebon", "30 hari", "Aktif"],
        ["SUP-014", "PT Marga Nusantara", "Bandung", "14 hari", "Aktif"],
        ["SUP-028", "Apex Jaya", "Cirebon", "7 hari", "Perlu tinjau"],
    ]),
    "Master Customer": ("Pelanggan", ["Kode", "Nama pelanggan", "Wilayah", "Termin", "Piutang"], [
        ["CUS-00011", "Ani TK", "C3", "14 hari", "Rp 1.912.000"],
        ["CUS-00016", "Maun TK", "C5", "7 hari", "Rp 2.498.500"],
        ["CUS-00018", "Lestari TK", "C1", "14 hari", "Rp 1.460.500"],
    ]),
    "Master Sales": ("Tim sales", ["Kode", "Nama", "Wilayah", "Target", "Status"], [
        ["S-01", "Masrukin", "Cirebon Barat", "Rp 120 jt", "Aktif"],
        ["S-02", "Tohirin", "Cirebon Timur", "Rp 110 jt", "Aktif"],
        ["S-03", "Nofal", "Kuningan", "Rp 95 jt", "Aktif"],
    ]),
    "Persediaan": ("Persediaan", ["SKU", "Produk", "Batch", "Stok", "Kedaluwarsa"], [
        ["000102", "Adem Sari", "AS2607", "396 HGR", "12/2027"],
        ["002847", "Bodrex Extra", "BE2608", "626 BOX", "08/2028"],
        ["000118", "Antimo", "AT2606", "670 LSN", "06/2028"],
    ]),
    "Stok Opname": ("Stok opname", ["Produk", "Sistem", "Fisik", "Selisih", "Status"], [
        ["Neozep F", "101", "90", "-11", "Perlu setujui"],
        ["Hansaplast", "248", "243", "-5", "Perlu setujui"],
        ["Fatigon Spirit", "0", "9", "+9", "Dihitung"],
    ]),
    "Analisis Harga": ("Analisis harga", ["Produk", "HPP", "Harga tunai", "Harga kredit", "Margin"], [
        ["Adem Sari", "Rp 50.416", "Rp 49.000", "Rp 49.500", "-1,8%"],
        ["Amplop 3/4 AM", "Rp 73.040", "Rp 76.000", "Rp 77.500", "4,1%"],
        ["Antimo", "Rp 50.183", "Rp 52.000", "Rp 54.000", "7,6%"],
    ]),
    "Harga Jual": ("Daftar harga jual", ["Produk", "Satuan", "Tunai", "Kredit", "Berlaku"], [
        ["Adem Sari", "HGR", "Rp 49.000", "Rp 49.500", "Aktif"],
        ["Bodrex Extra", "BOX", "Rp 54.000", "Rp 56.000", "Aktif"],
        ["Antimo", "LSN", "Rp 52.000", "Rp 54.000", "Aktif"],
    ]),
    "Ekspor Data": ("Ekspor data", ["Dataset", "Periode", "Baris", "Format", "Status"], [
        ["Stok dan harga", "5 Agu 2026", "626", "XLSX", "Siap"],
        ["Riwayat harga", "Jul-Agu 2026", "2.875", "XLSX", "Siap"],
        ["Batch-expiry", "Semua aktif", "1.204", "CSV", "Siap"],
    ]),
    "Laporan Persediaan": ("Laporan persediaan", ["Produk", "Awal", "Masuk", "Keluar", "Akhir"], [
        ["Adem Sari", "24", "1.878", "1.897", "5"],
        ["Amplox", "0", "164", "146", "18"],
        ["Antimo", "0", "2.952", "2.281", "670"],
    ]),
    "Master Harga": ("Buku harga", ["Nama buku", "Cakupan", "Berlaku", "Versi", "Status"], [
        ["Harga umum", "Semua pelanggan", "1 Agu 2026", "v12", "Disetujui"],
        ["Harga C1", "Wilayah C1", "1 Agu 2026", "v8", "Disetujui"],
        ["Harga promo", "Pelanggan pilihan", "5 Agu 2026", "v2", "Draft"],
    ]),
    "Master Harga Beli": ("Harga beli supplier", ["Supplier", "Produk", "Harga", "Tanggal", "Status"], [
        ["PT Marga Nusantara", "Napacin", "Rp 54.844", "23/10/2025", "Historis"],
        ["Praba Anton", "Antimo", "Rp 8.600", "04/03/2018", "Historis"],
        ["Aman Kiyong", "Paramex", "Rp 103.459", "05/08/2025", "Aktif"],
    ]),
    "Master Harga Jual": ("Harga khusus pelanggan", ["Pelanggan", "Produk", "Harga", "Tanggal", "Status"], [
        ["Tata", "Adem Sari", "Rp 53.000", "30/07/2026", "Aktif"],
        ["Tata", "Bodrex", "Rp 102.000", "30/07/2026", "Aktif"],
        ["Hj. Siti TK", "Antimo", "Rp 54.000", "04/08/2026", "Aktif"],
    ]),
    "Pembelian": ("Pembelian dan penerimaan", ["Faktur", "Supplier", "Produk", "Jumlah", "Status"], [
        ["PO-2608-012", "Wi Seng", "Bodrex Extra", "Rp 24.720.000", "Diterima"],
        ["PO-2608-013", "Agung Indah", "Tiger", "Rp 1.000.000", "Diterima"],
        ["PO-2608-014", "PD Zaki", "Signatur RK", "Rp 777.000", "Draft"],
    ]),
    "Hutang Dagang": ("Hutang supplier", ["Supplier", "Faktur", "Jatuh tempo", "Sisa", "Status"], [
        ["PT Bina San Prima", "060624", "27/06/2024", "Rp 5.458.760", "Jatuh tempo"],
        ["PT Dos Ni Roha", "110324", "19/04/2024", "Rp 717.948", "Jatuh tempo"],
        ["Wi Seng", "010826", "15/08/2026", "Rp 24.720.000", "Terbuka"],
    ]),
    "Pembayaran Hutang": ("Pembayaran hutang", ["Tanggal", "Supplier", "Faktur", "Metode", "Jumlah"], [
        ["05/08/2026", "Wi Seng", "010826", "Transfer", "Rp 10.000.000"],
        ["05/08/2026", "Agung Indah", "010826", "Tunai", "Rp 1.000.000"],
        ["04/08/2026", "PD Zaki", "040826", "Transfer", "Rp 1.621.000"],
    ]),
    "Riwayat Pembayaran Hutang": ("Riwayat pembayaran hutang", ["Supplier", "Faktur", "Tanggal bayar", "Jumlah", "Status"], [
        ["PT Bina San Prima", "030124", "30/01/2024", "Rp 3.100.229", "Terposting"],
        ["PT Bina San Prima", "060624", "27/06/2024", "Rp 5.458.760", "Terposting"],
        ["PT Dos Ni Roha", "110324", "19/04/2024", "Rp 717.948", "Terposting"],
    ]),
    "Laporan Pembayaran Hutang": ("Laporan pembayaran hutang", ["Supplier", "Faktur", "Tanggal", "Jumlah", "Bukti"], [
        ["PT Bina San Prima", "060624", "27/06/2024", "Rp 5.458.760", "PV-2406-018"],
        ["PT Dos Ni Roha", "110324", "19/04/2024", "Rp 717.948", "PV-2404-011"],
        ["Wi Seng", "010826", "05/08/2026", "Rp 10.000.000", "PV-2608-006"],
    ]),
    "Analisis Hutang": ("Analisis umur hutang", ["Supplier", "Belum jatuh tempo", "1-30", "31-60", ">90"], [
        ["Wi Seng", "Rp 14,7 jt", "Rp 10 jt", "-", "-"],
        ["PT Bina San Prima", "-", "-", "Rp 5,4 jt", "Rp 42 jt"],
        ["PT Dos Ni Roha", "-", "-", "Rp 717 rb", "Rp 8,2 jt"],
    ]),
    "Cetak Pembelian": ("Dokumen pembelian", ["Nomor", "Supplier", "Tanggal", "Total", "Versi"], [
        ["PI-2608-012", "Wi Seng", "01/08/2026", "Rp 24.720.000", "Asli"],
        ["PI-2608-013", "Agung Indah", "01/08/2026", "Rp 1.000.000", "Asli"],
        ["PI-2608-014", "PD Zaki", "02/08/2026", "Rp 1.621.000", "Cetak ulang"],
    ]),
    "Laporan Pembelian": ("Laporan pembelian", ["Supplier", "Faktur", "Tanggal", "Item", "Total"], [
        ["Wi Seng", "010826", "01/08/2026", "Bodrex Extra", "Rp 24.720.000"],
        ["Agung Indah", "010826", "01/08/2026", "Tiger", "Rp 1.000.000"],
        ["H. Memen", "020826", "02/08/2026", "King Kong", "Rp 2.600.000"],
    ]),
    "Penjualan": ("Order dan penjualan", ["Faktur", "Sales", "Pelanggan", "Total", "Status"], [
        ["2608-000530", "Masrukin", "Ani TK", "Rp 1.912.000", "Terkirim"],
        ["2608-000030", "Masrukin", "Maun TK", "Rp 587.500", "Terkirim"],
        ["2608-000018", "Tohirin", "Lestari TK", "Rp 364.000", "Draft"],
    ]),
    "Piutang Dagang": ("Piutang pelanggan", ["Pelanggan", "Faktur", "Jatuh tempo", "Sisa", "Status"], [
        ["Ani TK", "2607-000290", "02/08/2026", "Rp 266.500", "Terlambat"],
        ["Maun TK", "2608-000030", "06/08/2026", "Rp 587.500", "Terbuka"],
        ["Lestari TK", "2608-000018", "11/08/2026", "Rp 364.000", "Terbuka"],
    ]),
    "Penerimaan Piutang": ("Penerimaan piutang", ["Tanggal", "Pelanggan", "Faktur", "Metode", "Jumlah"], [
        ["05/08/2026", "Ani TK", "2607-000290", "Transfer", "Rp 266.500"],
        ["05/08/2026", "Maun TK", "2608-000030", "Tunai", "Rp 587.500"],
        ["04/08/2026", "Lestari TK", "2608-000018", "Giro", "Rp 364.000"],
    ]),
    "Riwayat Penerimaan Piutang": ("Riwayat penerimaan", ["Pelanggan", "Faktur", "Tanggal", "Jumlah", "Status"], [
        ["Ani TK", "2606-000343", "11/07/2026", "Rp 536.000", "Terposting"],
        ["Maun TK", "2607-000070", "08/07/2026", "Rp 852.500", "Terposting"],
        ["Lestari TK", "2606-000509", "15/07/2026", "Rp 347.000", "Terposting"],
    ]),
    "Laporan Penerimaan Piutang": ("Laporan penerimaan piutang", ["Pelanggan", "Faktur", "Tanggal", "Jumlah", "Bukti"], [
        ["Ani TK", "2607-000290", "05/08/2026", "Rp 266.500", "RV-2608-022"],
        ["Maun TK", "2608-000030", "05/08/2026", "Rp 587.500", "RV-2608-023"],
        ["Lestari TK", "2608-000018", "04/08/2026", "Rp 364.000", "RV-2608-019"],
    ]),
    "Analisis Piutang": ("Analisis piutang", ["Dimensi", "Belum jatuh tempo", "1-30", "31-60", ">90"], [
        ["Masrukin", "Rp 12,4 jt", "Rp 8,2 jt", "Rp 3,1 jt", "Rp 1,7 jt"],
        ["Tohirin", "Rp 9,8 jt", "Rp 5,4 jt", "Rp 2,8 jt", "Rp 900 rb"],
        ["Wilayah C1", "Rp 7,2 jt", "Rp 3,9 jt", "Rp 1,2 jt", "Rp 640 rb"],
    ]),
    "Distribusi Nota Sales": ("Serah-terima nota", ["Sales", "Pelanggan", "Faktur", "Nilai", "Status"], [
        ["Masrukin", "Ani TK", "2607-000530", "Rp 358.500", "Dibawa"],
        ["Masrukin", "Maun TK", "2608-000030", "Rp 587.500", "Dibawa"],
        ["Tohirin", "Lestari TK", "2608-000018", "Rp 364.000", "Dikembalikan"],
    ]),
    "Nota Penjualan": ("Nota penjualan", ["Pelanggan", "Faktur", "Tanggal", "Tagihan", "Checklist"], [
        ["Ani TK", "2607-000530", "31/07/2026", "Rp 358.500", "Belum"],
        ["Maun TK", "2608-000030", "04/08/2026", "Rp 587.500", "Belum"],
        ["Lestari TK", "2608-000018", "04/08/2026", "Rp 364.000", "Selesai"],
    ]),
    "Laporan Piutang": ("Laporan piutang", ["Pelanggan", "Faktur", "Tanggal", "Sisa", "Umur"], [
        ["Ani TK", "2607-000290", "31/07/2026", "Rp 266.500", "5 hari"],
        ["Maun TK", "2608-000030", "04/08/2026", "Rp 587.500", "1 hari"],
        ["Lestari TK", "2608-000018", "04/08/2026", "Rp 364.000", "1 hari"],
    ]),
    "Kas/Jurnal": ("Jurnal umum", ["Bukti", "Tanggal", "Akun debit", "Akun kredit", "Jumlah"], [
        ["JV-2608-001", "05/08/2026", "Kas", "Pendapatan lain", "Rp 500.000"],
        ["JV-2608-002", "05/08/2026", "Biaya angkut", "Kas", "Rp 125.000"],
        ["JV-2608-003", "05/08/2026", "Piutang", "Penjualan", "Rp 1.912.000"],
    ]),
    "Master Perkiraan": ("Daftar perkiraan", ["Kode", "Nama akun", "Kelompok", "Saldo", "Status"], [
        ["102", "Piutang dagang", "Aktiva", "Rp 99.675.925", "Aktif"],
        ["202", "Hutang supplier", "Kewajiban", "Rp 11.917.398", "Aktif"],
        ["500", "Biaya operasional", "Biaya", "Rp 8.425.000", "Aktif"],
    ]),
    "Laporan Laba Rugi": ("Laba rugi", ["Komponen", "Bulan ini", "Bulan lalu", "Perubahan", "Status"], [
        ["Penjualan bersih", "Rp 428,6 jt", "Rp 396,2 jt", "+8,2%", "Baik"],
        ["HPP", "Rp 367,2 jt", "Rp 341,4 jt", "+7,6%", "Pantau"],
        ["Laba kotor", "Rp 61,4 jt", "Rp 54,8 jt", "+12,0%", "Baik"],
    ]),
    "Laba Rugi Kotor": ("Laba kotor", ["Sales", "Omzet", "HPP", "Laba", "Margin"], [
        ["Masrukin", "Rp 126 jt", "Rp 107 jt", "Rp 19 jt", "15,1%"],
        ["Tohirin", "Rp 109 jt", "Rp 94 jt", "Rp 15 jt", "13,8%"],
        ["Nofal", "Rp 98 jt", "Rp 83 jt", "Rp 15 jt", "15,3%"],
    ]),
}


def slugify(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", "-", value.lower()).strip("-")


def words(text: str) -> int:
    return len(re.findall(r"\b[\w-]+\b", text, flags=re.UNICODE))


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        Path(os.environ.get("WINDIR", "C:/Windows")) / "Fonts" / ("arialbd.ttf" if bold else "arial.ttf"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def fit(draw: ImageDraw.ImageDraw, text: str, box: tuple[int, int, int, int], size: int, color: str, bold: bool = False) -> None:
    x1, y1, x2, y2 = box
    selected = font(size, bold)
    line = ""
    lines: list[str] = []
    for token in text.split():
        trial = f"{line} {token}".strip()
        if draw.textbbox((0, 0), trial, font=selected)[2] <= x2 - x1:
            line = trial
        else:
            lines.append(line)
            line = token
    if line:
        lines.append(line)
    line_height = size + 7
    for index, value in enumerate(lines[: max(1, (y2 - y1) // line_height)]):
        draw.text((x1, y1 + index * line_height), value, font=selected, fill=color)


def status_label(value: str) -> str:
    return {"OPERATIONAL": "Operasional", "READ_ONLY": "Baca saja", "CONTRACT_ONLY": "Kontrak"}.get(value, value)


def parse_parity() -> dict[int, tuple[str, str]]:
    values: dict[int, tuple[str, str]] = {}
    pattern = re.compile(r"item\((\d+),.*?'(OPERATIONAL|READ_ONLY|CONTRACT_ONLY)',\s*'(OPERATIONAL|READ_ONLY|CONTRACT_ONLY)'\)")
    for line in PARITY_SOURCE.read_text(encoding="utf-8").splitlines():
        match = pattern.search(line)
        if match:
            values[int(match.group(1))] = (match.group(2), match.group(3))
    if len(values) != 48:
        raise ValueError(f"Status paritas terbaca {len(values)}, seharusnya 48")
    return values


def read_matrix() -> list[dict[str, str]]:
    with MATRIX.open(encoding="utf-8-sig", newline="") as handle:
        rows = list(csv.DictReader(handle))
    if len(rows) != 48:
        raise ValueError(f"Matriks berisi {len(rows)} layar, seharusnya 48")
    return rows


def extract_legacy_chapters() -> tuple[dict[int, list[str]], dict[int, str]]:
    document = Document(LEGACY_DOCX)
    chapters: dict[int, list[str]] = {}
    media: dict[int, str] = {}
    current: int | None = None
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        match = re.match(r"Layar\s+(\d+)\s+-", text)
        if match:
            current = int(match.group(1))
            chapters[current] = []
            continue
        if current is None:
            continue
        if text and not text.startswith("Jumlah kata uraian layar"):
            chapters[current].append(text)
        for blip in paragraph._p.xpath(".//a:blip"):
            relationship_id = blip.get(qn("r:embed"))
            media[current] = Path(document.part.rels[relationship_id].target_ref).name
    if len(chapters) != 48 or len(media) != 48:
        raise ValueError(f"Manual legacy terbaca {len(chapters)} bab dan {len(media)} relasi gambar")
    return chapters, media


def extract_legacy_media(target: Path) -> None:
    target.mkdir(parents=True, exist_ok=True)
    with ZipFile(LEGACY_DOCX) as archive:
        for name in archive.namelist():
            if name.startswith("word/media/"):
                (target / Path(name).name).write_bytes(archive.read(name))


def module_view(module: str) -> tuple[str, list[str], list[list[str]]]:
    if module in DOMAIN_CONFIG:
        return DOMAIN_CONFIG[module]
    if "Pembelian" in module or "Hutang" in module:
        return DOMAIN_CONFIG["Hutang Dagang"]
    if "Piutang" in module:
        return DOMAIN_CONFIG["Piutang Dagang"]
    if "Harga" in module:
        return DOMAIN_CONFIG["Analisis Harga"]
    if "Laba" in module:
        return DOMAIN_CONFIG["Laporan Laba Rugi"]
    return DOMAIN_CONFIG["Persediaan"]


def generate_modern_image(row: dict[str, str], web: str, flutter: str) -> Path:
    number = int(row["no"])
    title, headers, rows = module_view(row["module"])
    path = MODERN_IMAGES / f"{number:02d}-{slugify(row['legacy_screen'])}.png"
    canvas = Image.new("RGB", (1600, 900), "#F4F7FB")
    draw = ImageDraw.Draw(canvas)

    draw.rectangle((0, 0, 236, 900), fill="#101A2E")
    draw.rounded_rectangle((30, 28, 72, 70), radius=9, fill="#0F766E")
    draw.text((43, 38), "eB", font=font(15, True), fill="white")
    draw.text((88, 38), "CMN Inventory", font=font(17, True), fill="white")
    menu = ["Beranda", "Master data", "Produk & stok", "Pembelian", "Penjualan", "Keuangan", "Laporan"]
    selected = min(6, max(1, number // 8))
    for index, label in enumerate(menu):
        y = 110 + index * 58
        if index == selected:
            draw.rounded_rectangle((20, y - 8, 216, y + 34), radius=8, fill="#1E3A5F")
        draw.ellipse((36, y + 2, 48, y + 14), fill="#5EEAD4" if index == selected else "#94A3B8")
        draw.text((62, y), label, font=font(15, index == selected), fill="white" if index == selected else "#CBD5E1")

    draw.rectangle((236, 0, 1600, 78), fill="white", outline="#DCE3EC")
    draw.text((270, 27), "Caruban Medika Nusantara", font=font(18, True), fill="#0F172A")
    draw.rounded_rectangle((1250, 18, 1395, 58), radius=20, fill="#ECFDF5", outline="#A7F3D0")
    draw.text((1272, 29), "Data tersinkron", font=font(13, True), fill="#047857")
    draw.rounded_rectangle((1410, 18, 1565, 58), radius=20, fill="#F8FAFC", outline="#CBD5E1")
    draw.text((1432, 29), "Muklis / Pemilik", font=font(13, True), fill="#334155")

    draw.text((274, 112), f"PADANAN LAYAR LAMA {number:02d}", font=font(13, True), fill="#0F766E")
    fit(draw, row["legacy_screen"], (274, 142, 1040, 220), 31, "#0F172A", True)
    draw.text((274, 220), f"Ruang kerja: {title}  /  Modul ERP: {row['module']}", font=font(15), fill="#64748B")

    web_fill = "#DCFCE7" if web == "OPERATIONAL" else "#FEF3C7"
    web_color = "#166534" if web == "OPERATIONAL" else "#92400E"
    flutter_fill = "#DCFCE7" if flutter == "OPERATIONAL" else "#E2E8F0"
    flutter_color = "#166534" if flutter == "OPERATIONAL" else "#475569"
    draw.rounded_rectangle((1160, 120, 1340, 158), radius=18, fill=web_fill)
    draw.text((1182, 131), f"Web: {status_label(web)}", font=font(13, True), fill=web_color)
    draw.rounded_rectangle((1350, 120, 1565, 158), radius=18, fill=flutter_fill)
    draw.text((1371, 131), f"Flutter: {status_label(flutter)}", font=font(13, True), fill=flutter_color)

    controls = ["Cari nama, kode, atau nomor dokumen", "Periode", "Status"]
    x = 274
    widths = [520, 180, 170]
    for label, width in zip(controls, widths):
        draw.rounded_rectangle((x, 268, x + width, 320), radius=8, fill="white", outline="#CBD5E1", width=2)
        draw.text((x + 18, 284), label, font=font(14), fill="#64748B")
        x += width + 16
    draw.rounded_rectangle((1328, 268, 1565, 320), radius=8, fill="#0F766E")
    draw.text((1392, 284), "+ Tambah / Proses", font=font(14, True), fill="white")

    draw.rounded_rectangle((274, 346, 1565, 734), radius=12, fill="white", outline="#D8E0EA", width=2)
    col_width = 1240 // len(headers)
    for index, header in enumerate(headers):
        x = 298 + index * col_width
        draw.text((x, 372), header, font=font(14, True), fill="#475569")
    draw.line((296, 404, 1538, 404), fill="#CBD5E1", width=2)
    for row_index, values in enumerate(rows):
        y = 430 + row_index * 72
        if row_index == 0:
            draw.rounded_rectangle((290, y - 12, 1548, y + 42), radius=7, fill="#ECFDF5")
        for col_index, value in enumerate(values):
            x = 298 + col_index * col_width
            fit(draw, value, (x, y, x + col_width - 15, y + 38), 14, "#0F172A", col_index == 1)
        draw.line((296, y + 52, 1538, y + 52), fill="#E2E8F0")

    draw.rounded_rectangle((274, 762, 900, 864), radius=10, fill="#EFF6FF", outline="#BFDBFE")
    draw.text((296, 782), "Perubahan utama", font=font(14, True), fill="#1D4ED8")
    fit(draw, "Data ditampilkan terstruktur, dapat dicari, memiliki status, jejak audit, dan hak akses per peran.", (296, 810, 876, 858), 12, "#334155")
    draw.rounded_rectangle((924, 762, 1565, 864), radius=10, fill="#FFF7ED", outline="#FED7AA")
    draw.text((946, 782), "Kontrol sebelum selesai", font=font(14, True), fill="#C2410C")
    fit(draw, shorten(row["primary_risk"], width=120, placeholder="..."), (946, 810, 1540, 858), 12, "#334155")
    canvas.save(path, optimize=True)
    return path


def transition_sections(row: dict[str, str], web: str, flutter: str) -> list[tuple[str, list[str]]]:
    title = row["legacy_screen"]
    module = row["module"]
    business = row["business_focus"].rstrip(".")
    parity = row["parity_rule"].rstrip(".")
    risk = row["primary_risk"].rstrip(".")
    acceptance = row["minimum_acceptance"].rstrip(".")
    return [
        ("Makna perbandingan layar lama dan layar baru", [
            f"Layar lama {title} ditempatkan di samping ilustrasi padanan baru agar pengguna dapat mengenali pekerjaan yang sama tanpa harus menghafal tata letak baru. Tujuan bisnisnya tetap {business}. Perubahan utama terletak pada cara sistem mengelompokkan informasi, memberi status, mencatat jejak audit, dan membatasi tindakan sesuai peran. Pada aplikasi lama, tombol dan tabel sering berada dalam satu jendela padat. Pada sistem baru, daftar, detail, tindakan, persetujuan, dan laporan dipisahkan secara lebih jelas supaya kesalahan klik berkurang dan pengguna memahami akibat setiap tindakan.",
            f"Padanan baru berada dalam modul {module}. Status yang dapat dibuktikan pada saat manual ini dibuat adalah Web {status_label(web).lower()} dan Flutter {status_label(flutter).lower()}. Status tersebut penting: operasional berarti alur transaksi telah tersedia pada permukaan terkait; baca saja berarti pengguna dapat meninjau data tetapi tindakan perubahan belum dilakukan dari permukaan itu; kontrak berarti kebutuhan dan API telah dipetakan tetapi layar Flutter lengkap belum boleh dianggap tersedia. Manual sengaja menampilkan status ini agar pelatihan tidak menjanjikan fungsi yang belum dibuktikan.",
        ]),
        ("Cara menemukan fungsi yang dahulu berada pada jendela legacy", [
            f"Pengguna lama tidak lagi mencari tombol berdasarkan posisi ikon. Mulailah dari nama proses {module}, lalu gunakan pencarian global atau menu sisi kiri. Ketik nama pelanggan, supplier, produk, nomor faktur, atau istilah {title} sesuai kebutuhan. Setelah daftar muncul, gunakan filter periode dan status. Klik satu baris untuk membuka rincian. Tombol tindakan hanya muncul bila status dokumen dan hak akses mengizinkan. Pola ini konsisten di layar Web, Desktop Windows, dan bagian Flutter yang sudah operasional sehingga pengguna tidak perlu mempelajari tata letak berbeda untuk setiap transaksi.",
            f"Aturan bisnis lama tetap dihormati. {parity}. Namun penerapannya tidak lagi bergantung pada warna baris atau urutan indeks DBF semata. Sistem baru menyimpan identitas data, waktu perubahan, pengguna, status dokumen, dan relasi ke transaksi lain. Bila data berasal dari migrasi, kode legacy dipertahankan sebagai referensi rekonsiliasi. Pengguna sebaiknya mencari dengan kode lama terlebih dahulu ketika mencocokkan arsip, lalu memastikan nama dan pihak terkait sebelum melakukan tindakan.",
        ]),
        ("Langkah kerja transisi untuk pengguna nonteknis", [
            f"Pertama, pastikan nama tenant Caruban Medika Nusantara, cabang, periode, dan identitas pengguna pada bagian atas layar sudah benar. Kedua, buka modul {module} dan pilih fungsi yang paling mendekati istilah lama {title}. Ketiga, gunakan filter untuk mempersempit data sebelum membuka rincian. Keempat, bandingkan nomor referensi, tanggal, pihak, produk, jumlah, harga, dan status dengan dokumen sumber. Kelima, lakukan tindakan yang tersedia, misalnya menyimpan draft, mengajukan persetujuan, memposting, menerima pembayaran, atau membuat snapshot laporan. Keenam, baca pesan hasil dan simpan nomor bukti. Ketujuh, periksa kembali daftar untuk memastikan status berubah sesuai harapan.",
            f"Bila tombol tidak terlihat, jangan mencoba masuk dengan akun orang lain. Periksa peran pengguna, status dokumen, periode akuntansi, serta status sinkronisasi. Bila Flutter menunjukkan baca saja atau kontrak, lanjutkan tindakan melalui Web sesuai prosedur kantor. Bila koneksi terputus ketika fungsi offline tersedia, tunggu status antrean dan jangan membuat transaksi kedua untuk pekerjaan yang sama. Setelah jaringan kembali, lakukan sinkronisasi dan cocokkan nomor referensi server. Langkah sederhana ini mencegah transaksi ganda dan menjaga pertanggungjawaban pengguna.",
        ]),
        ("Kontrol data, risiko, dan bukti penyelesaian", [
            f"Risiko utama pada proses ini adalah {risk}. Sistem baru membantu melalui validasi, status, audit log, serta pemisahan kewenangan, tetapi pengguna tetap bertanggung jawab membaca data. Sebelum menyimpan, periksa satuan, nilai rupiah, tanggal, pihak, gudang, sales, dan lampiran. Untuk transaksi stok atau keuangan, pastikan total rincian sama dengan total dokumen. Untuk perubahan master, pastikan tidak membuat duplikat. Untuk laporan, simpan parameter, waktu pembentukan, versi, dan siapa yang meminta cetak ulang.",
            f"Bukti minimum penyelesaian adalah nomor dokumen atau kode referensi, status akhir, waktu proses, nama pengguna, serta hasil rekonsiliasi. Kriteria penerimaannya: {acceptance}. Apabila salah satu bukti belum tersedia, pekerjaan belum boleh dianggap selesai walaupun layar telah ditutup. Catat masalah pada log dukungan dengan nomor referensi dan tangkapan layar yang tidak memuat kata sandi atau data sensitif berlebihan.",
        ]),
        ("Pembagian penggunaan Web, Windows, dan Android", [
            f"Web menjadi ruang kerja utama untuk proses {module} yang memerlukan tabel lebar, pemeriksaan lintas dokumen, persetujuan, atau ekspor. Aplikasi Windows Flutter mengikuti pola navigasi yang lebih ringkas untuk operasi yang telah disediakan dan dapat dipakai pada komputer operasional. Android Flutter diprioritaskan untuk pekerjaan lapangan seperti katalog, order, penerimaan piutang, serah-terima nota, dan pemantauan. Pada layar yang belum operasional di Flutter, pengguna tetap dapat melihat status paritas dan diarahkan ke Web, bukan diberi tombol semu.",
            f"Tidak semua pekerjaan harus dilakukan di semua perangkat. Pemilihan perangkat mengikuti risiko dan konteks. Entri cepat saat kunjungan cocok untuk Android; pemeriksaan rinci dan laporan cocok untuk Web atau Windows; persetujuan bernilai besar sebaiknya dilakukan pada perangkat pribadi yang aman. Pada semua perangkat, pengguna harus keluar setelah selesai, menjaga PIN atau kata sandi, dan memastikan sinkronisasi tidak tertunda sebelum mengganti atau menghapus aplikasi.",
        ]),
    ]


def set_cell_shading(cell, fill_value: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_value)
    cell._tc.get_or_add_tcPr().append(shading)


def set_repeat_table_header(row) -> None:
    table_properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    table_properties.append(repeat)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def configure(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.12
    for style_name, size, color in [("Title", 31, INK), ("Heading 1", 23, INK), ("Heading 2", 15, TEAL), ("Heading 3", 12, INK)]:
        style = styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
    header = section.header.paragraphs[0]
    header.text = "eBisnis Inventory / Sales  |  Panduan transisi 48 layar"
    header.style = styles["Caption"]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Caruban Medika Nusantara  |  Halaman ")
    add_page_number(footer)


def add_cover(document: Document, modern_cover: Path) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("PANDUAN TRANSISI OPERASIONAL")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = TEAL
    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("48 Layar Lama ke Sistem Baru\neBisnis Inventory / Sales")
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Perbandingan visual, langkah kerja, kontrol, dan status paritas Web, Windows, serta Android Flutter").italic = True
    image_paragraph = document.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.add_run().add_picture(str(modern_cover), width=Inches(6.65))
    table = document.add_table(rows=2, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    values = [("TENANT CONTOH", "Caruban Medika Nusantara"), ("CAKUPAN", "48 pasangan layar"), ("VERSI", "1.0"), ("STATUS", "Manual transisi"), ("SUMBER", "Manual legacy + matriks paritas"), ("DIPERBARUI", "5 Agustus 2026")]
    for index, (label, value) in enumerate(values):
        cell = table.cell(index // 3, index % 3)
        set_cell_shading(cell, "F1F5F9")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{label}\n")
        r.bold = True
        r.font.size = Pt(8)
        p.add_run(value).font.size = Pt(9)
    document.add_page_break()


def add_front_matter(document: Document, rows: list[dict[str, str]], parity: dict[int, tuple[str, str]]) -> None:
    document.add_heading("Cara memakai panduan transisi", level=1)
    for text in [
        "Panduan ini ditujukan bagi pengguna yang telah terbiasa dengan INVENTORY CONTROL berbasis DBF. Setiap bab memperlihatkan layar lama terlebih dahulu, kemudian ilustrasi padanan pada eBisnis. Gunakan gambar lama untuk mengenali pekerjaan, lalu ikuti nama modul, filter, status, dan tindakan pada gambar baru.",
        "Ilustrasi sistem baru memakai data pelatihan Caruban Medika Nusantara. Gambar tidak menyalin data produksi dan tidak menggantikan hak akses. Status Operasional, Baca saja, dan Kontrak diambil dari katalog paritas pada saat dokumen dibuat. Status dapat meningkat setelah implementasi, pengujian, dan UAT, tetapi tidak boleh diasumsikan lebih tinggi dari yang tertulis.",
        "Setiap bab mempunyai lebih dari 1.500 kata. Uraian lama dipertahankan karena memuat pengetahuan proses dan risiko yang masih relevan. Bagian transisi menjelaskan lokasi baru, perubahan istilah, pembagian perangkat, dan bukti penyelesaian.",
    ]:
        document.add_paragraph(text)
    document.add_heading("Indeks 48 layar", level=1)
    table = document.add_table(rows=1, cols=5)
    table.style = "Light Shading Accent 1"
    headers = ["No.", "Layar lama", "Modul baru", "Web", "Flutter"]
    for index, header in enumerate(headers):
        table.cell(0, index).text = header
        set_cell_shading(table.cell(0, index), "0F766E")
        for run in table.cell(0, index).paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True
    set_repeat_table_header(table.rows[0])
    for row in rows:
        number = int(row["no"])
        cells = table.add_row().cells
        values = [str(number), row["legacy_screen"], row["module"], status_label(parity[number][0]), status_label(parity[number][1])]
        for index, value in enumerate(values):
            cells[index].text = value
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    document.add_page_break()


def add_picture(document: Document, path: Path, caption_text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    inline = paragraph.add_run().add_picture(str(path), width=Inches(6.75))
    inline._inline.docPr.set("descr", caption_text)
    caption = document.add_paragraph(style="Caption")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.add_run(caption_text).italic = True


def add_chapter(document: Document, row: dict[str, str], legacy_image: Path, modern_image: Path, legacy_paragraphs: list[str], web: str, flutter: str) -> int:
    number = int(row["no"])
    document.add_heading(f"{number:02d}. {row['legacy_screen']}", level=1)
    lead = document.add_paragraph()
    run = lead.add_run(f"Modul padanan: {row['module']}  |  Web: {status_label(web)}  |  Flutter: {status_label(flutter)}")
    run.bold = True
    run.font.color.rgb = TEAL
    document.add_heading("Tampilan aplikasi lama", level=2)
    add_picture(document, legacy_image, f"Gambar {number}A. Layar legacy {row['legacy_screen']} dari manual INVENTORY CONTROL.")
    document.add_heading("Ilustrasi padanan pada sistem baru", level=2)
    add_picture(document, modern_image, f"Gambar {number}B. Ilustrasi padanan eBisnis untuk {row['legacy_screen']}; data hanya contoh pelatihan.")

    comparison = document.add_table(rows=1, cols=2)
    comparison.alignment = WD_TABLE_ALIGNMENT.CENTER
    comparison.style = "Light Shading Accent 1"
    comparison.cell(0, 0).text = "Yang dipertahankan"
    comparison.cell(0, 1).text = "Yang berubah"
    for cell in comparison.rows[0].cells:
        set_cell_shading(cell, "0F766E")
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True
    cells = comparison.add_row().cells
    cells[0].text = row["business_focus"]
    cells[1].text = "Navigasi berbasis modul, pencarian, filter, status dokumen, hak akses, jejak audit, dan sinkronisasi menggantikan ketergantungan pada posisi ikon serta indeks DBF."

    sections = transition_sections(row, web, flutter)
    for heading, paragraphs in sections:
        document.add_heading(heading, level=2)
        for text in paragraphs:
            document.add_paragraph(text)
    document.add_heading("Pengetahuan operasional dari manual lama yang tetap berlaku", level=2)
    for text in legacy_paragraphs:
        document.add_paragraph(text)
    all_text = " ".join(text for _, paragraphs in sections for text in paragraphs) + " " + " ".join(legacy_paragraphs)
    count = words(all_text)
    if count < 1500:
        raise ValueError(f"Layar {number} hanya {count} kata")
    check = document.add_paragraph()
    check.paragraph_format.keep_with_next = True
    check_run = check.add_run(f"Verifikasi bab: {count:,} kata penjelasan. Minimum 1.500 kata terpenuhi.")
    check_run.bold = True
    check_run.font.color.rgb = TEAL
    document.add_page_break()
    return count


def build() -> dict[str, object]:
    if not LEGACY_DOCX.exists() or not MATRIX.exists():
        raise FileNotFoundError("Paket dokumentasi legacy tidak ditemukan. Atur INVENTORY_LEGACY_PACKAGE bila lokasinya berbeda.")
    PUBLIC.mkdir(parents=True, exist_ok=True)
    MODERN_IMAGES.mkdir(parents=True, exist_ok=True)
    INDEX.parent.mkdir(parents=True, exist_ok=True)
    rows = read_matrix()
    parity = parse_parity()
    legacy_chapters, media_map = extract_legacy_chapters()
    counts: dict[str, int] = {}
    index_rows: list[dict[str, object]] = []

    with tempfile.TemporaryDirectory(prefix="ebisnis-legacy-manual-") as temp_name:
        temp = Path(temp_name)
        extract_legacy_media(temp)
        modern_paths: dict[int, Path] = {}
        for row in rows:
            number = int(row["no"])
            modern_paths[number] = generate_modern_image(row, *parity[number])

        document = Document()
        configure(document)
        add_cover(document, modern_paths[1])
        add_front_matter(document, rows, parity)
        for row in rows:
            number = int(row["no"])
            count = add_chapter(document, row, temp / media_map[number], modern_paths[number], legacy_chapters[number], *parity[number])
            counts[str(number)] = count
            index_rows.append({
                "screen": number,
                "legacyName": row["legacy_screen"],
                "module": row["module"],
                "web": parity[number][0],
                "flutter": parity[number][1],
                "wordCount": count,
                "modernImage": f"/panduan/inventory-sales/images/transisi/baru/{modern_paths[number].name}",
            })

        document.add_heading("Penutup dan penggunaan saat pelatihan", level=1)
        document.add_paragraph("Gunakan panduan ini sebagai pendamping pada masa transisi, bukan sebagai alasan mempertahankan kebiasaan yang berisiko. Pelatih sebaiknya memilih bab sesuai tugas peserta, memperagakan data contoh, meminta peserta mengulangi alur, lalu memeriksa bukti penyelesaian. Temuan yang berbeda dari manual harus dicatat sebagai bahan perbaikan sistem, materi, atau prosedur kerja.")
        properties = document.core_properties
        properties.title = "Panduan Transisi 48 Layar eBisnis Inventory / Sales"
        properties.subject = "Perbandingan layar legacy dan padanan sistem baru untuk pengguna nonteknis"
        properties.author = "eBisnis.id"
        properties.keywords = "inventory, sales, legacy, transisi, flutter, android, windows, web, Caruban Medika Nusantara"
        document.save(DOCX)

    INDEX.write_text(json.dumps({"meta": {"title": "Indeks Transisi 48 Layar", "version": "1.0", "updated": "5 Agustus 2026"}, "screens": index_rows}, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    REPORT.write_text(json.dumps({"minimumWordsPerScreen": 1500, "screens": 48, "counts": counts, "minimum": min(counts.values()), "maximum": max(counts.values()), "total": sum(counts.values())}, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    return {"docx": str(DOCX), "index": str(INDEX), "report": str(REPORT), "minimum": min(counts.values()), "total": sum(counts.values())}


if __name__ == "__main__":
    print(json.dumps(build(), ensure_ascii=False, indent=2))

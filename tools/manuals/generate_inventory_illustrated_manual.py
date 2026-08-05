from __future__ import annotations

import json
import re
from pathlib import Path
from textwrap import wrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
PUBLIC = ROOT / "apps/web/public/panduan/inventory-sales"
IMAGES = PUBLIC / "images/operasional"
CONTENT = ROOT / "apps/web/src/pages/inventory/inventory-illustrated-manual-content.json"
DOCX = PUBLIC / "Panduan-Operasional-Bergambar-eBisnis-Inventory-Sales.docx"
REPORT = ROOT / "docs/user-manual/inventory-illustrated-word-count.json"

INK = RGBColor(15, 23, 42)
TEAL = RGBColor(15, 118, 110)
MUTED = RGBColor(71, 85, 105)
PAGE_WIDTH_DXA = 9360


CHAPTERS = [
    {
        "id": "dashboard-pemilik",
        "title": "Dashboard Pemilik dan Ringkasan Kinerja",
        "image": "01-dashboard-pemilik.png",
        "platform": "Web dan Desktop Windows",
        "role": "Pemilik usaha dan manajemen",
        "objective": "memahami omzet, laba kotor, piutang, persediaan, risiko kedaluwarsa, dan kinerja setiap sales tanpa harus membuka transaksi satu per satu",
        "regions": ["periode laporan", "kartu indikator utama", "grafik tren penjualan", "peringkat sales", "peringatan stok", "daftar tindakan yang perlu diputuskan"],
        "steps": ["pastikan periode dan cabang sudah benar", "baca omzet bersama laba, bukan omzet sendirian", "bandingkan target dengan realisasi sales", "periksa piutang yang melewati jatuh tempo", "buka rincian batch yang mendekati kedaluwarsa", "catat keputusan dan penanggung jawab tindak lanjut"],
        "validations": ["angka kartu sama dengan rincian laporan", "periode tidak tercampur dengan bulan sebelumnya", "retur dan pembatalan sudah mengurangi omzet", "sales tanpa transaksi tetap terlihat", "peringatan stok mempunyai nama barang dan lokasi", "data terakhir tersinkron tercantum"],
        "risks": ["menganggap omzet sebagai laba", "membandingkan sales dengan wilayah yang tidak setara", "mengabaikan transaksi yang belum tersinkron", "mengambil keputusan hanya dari satu hari", "membaca stok tanpa memperhatikan batch", "membagikan dashboard kepada pihak yang tidak berwenang"],
        "scenario": "Muklis membuka dashboard pada awal pagi untuk melihat penjualan empat sales, mengecek piutang yang harus ditagih, lalu menentukan produk yang perlu diprioritaskan sebelum kunjungan lapangan dimulai.",
        "outcome": "Pemilik mendapatkan satu daftar keputusan harian yang dapat ditindaklanjuti, sementara sales menerima arahan yang jelas dan gudang memahami barang yang perlu disiapkan.",
        "stats": [("Omzet bulan ini", "Rp 428,6 jt"), ("Laba kotor", "Rp 61,4 jt"), ("Piutang jatuh tempo", "Rp 18,9 jt"), ("Batch berisiko", "12")],
        "table": ["Sales", "Omzet", "Tagihan", "Kunjungan", "Status"],
        "rows": [["Masrukin", "Rp 126 jt", "Rp 5,2 jt", "86", "Baik"], ["Tohirin", "Rp 109 jt", "Rp 4,8 jt", "79", "Pantau"], ["Nofal", "Rp 98 jt", "Rp 3,1 jt", "74", "Baik"], ["Agung", "Rp 95 jt", "Rp 5,8 jt", "71", "Tindak lanjut"]],
    },
    {
        "id": "sales-mobile",
        "title": "Kunjungan dan Order Sales melalui Android Flutter",
        "image": "02-sales-mobile.png",
        "platform": "Android Flutter",
        "role": "Sales lapangan",
        "objective": "menjalankan kunjungan pelanggan, memeriksa katalog dan stok, membuat order, serta menyimpan bukti kegiatan secara tertib meskipun koneksi internet tidak selalu stabil",
        "regions": ["identitas sales dan status sinkronisasi", "rute kunjungan", "profil pelanggan", "katalog dan pencarian produk", "keranjang order", "ringkasan pengiriman dan pembayaran"],
        "steps": ["masuk menggunakan akun pribadi", "sinkronkan data sebelum berangkat", "pilih pelanggan pada rute hari ini", "konfirmasi alamat dan batas kredit", "tambahkan produk beserta jumlahnya", "periksa harga, potongan, dan stok", "simpan order", "pastikan status terkirim ke server"],
        "validations": ["pelanggan sesuai wilayah sales", "harga berasal dari buku harga aktif", "jumlah tidak melebihi kebijakan stok", "potongan mempunyai dasar yang sah", "tanggal kirim dapat dipenuhi", "order mempunyai nomor referensi", "status offline atau online terlihat jelas"],
        "risks": ["memakai akun sales lain", "membuat order untuk pelanggan yang salah", "menjanjikan stok yang belum tersedia", "menghapus aplikasi sebelum sinkronisasi", "menyimpan kata sandi di catatan terbuka", "menganggap tombol simpan sama dengan berhasil tersinkron"],
        "scenario": "Masrukin mengunjungi toko pelanggan di wilayah Cirebon, memilih produk yang biasa dibeli, menyesuaikan jumlah dengan stok dan batas kredit, lalu mengirim order saat koneksi kembali tersedia.",
        "outcome": "Order lapangan tercatat dengan identitas sales, pelanggan, waktu, harga, dan status sinkronisasi sehingga admin tidak perlu mengetik ulang catatan kertas.",
        "stats": [("Kunjungan", "12/16"), ("Order hari ini", "9"), ("Nilai order", "Rp 18,4 jt"), ("Belum sinkron", "2")],
        "table": ["Pelanggan", "Wilayah", "Agenda", "Order", "Status"],
        "rows": [["Toko Sehat", "C1", "09.00", "Rp 2,8 jt", "Terkirim"], ["Apotek Melati", "C2", "10.30", "Rp 4,1 jt", "Terkirim"], ["Toko Maju", "C3", "13.00", "Rp 1,9 jt", "Offline"], ["Klinik Harmoni", "C4", "15.00", "-", "Berikutnya"]],
    },
    {
        "id": "produk-batch-harga",
        "title": "Produk, Batch, Kedaluwarsa, dan Harga Pelanggan",
        "image": "03-produk-batch-harga.png",
        "platform": "Web, Desktop Windows, dan Android Flutter",
        "role": "Admin master data, gudang, sales, dan pemilik",
        "objective": "menjaga satu sumber data produk yang lengkap, menelusuri batch dan kedaluwarsa, serta memastikan harga jual sesuai pelanggan dan kebijakan margin",
        "regions": ["kode dan nama produk", "satuan dan isi kemasan", "stok tersedia", "batch dan tanggal kedaluwarsa", "harga tunai dan kredit", "riwayat perubahan harga"],
        "steps": ["cari produk dengan kode atau nama", "periksa satuan penjualan", "lihat stok per gudang dan batch", "utamakan batch sesuai kebijakan FEFO", "pilih buku harga pelanggan", "periksa margin sebelum menyimpan", "catat alasan perubahan harga"],
        "validations": ["kode produk tidak ganda", "satuan dasar dan satuan jual konsisten", "batch mempunyai tanggal penerimaan", "produk kedaluwarsa tidak dapat dijual", "harga tidak berada di bawah batas tanpa persetujuan", "riwayat harga menyimpan pengguna dan waktu"],
        "risks": ["menjual satuan yang keliru", "mengabaikan tanggal kedaluwarsa", "mengubah harga lama tanpa tanggal berlaku", "mengandalkan stok total tanpa melihat batch", "memakai foto produk yang menyesatkan", "menghapus produk yang sudah dipakai transaksi"],
        "scenario": "Admin memperbarui harga obat yang baru diterima, gudang memeriksa batch, dan sales melihat harga khusus pelanggan yang berlaku tanpa mengubah data master dari ponsel.",
        "outcome": "Nama, satuan, stok, batch, dan harga tetap konsisten pada katalog publik, order sales, penerimaan gudang, invoice, dan laporan laba.",
        "stats": [("Produk aktif", "626"), ("Batch tersedia", "2.875"), ("Mendekati expiry", "24"), ("Harga perlu tinjau", "7")],
        "table": ["Produk", "Batch", "Expiry", "Stok", "Harga tunai"],
        "rows": [["Adem Sari", "AS-0826", "08/2027", "55", "Rp 49.000"], ["Amplop 3/4", "AM-0726", "07/2028", "34", "Rp 76.000"], ["Antimo", "AN-0626", "06/2027", "670", "Rp 52.000"], ["Bodrex", "BD-0526", "05/2027", "879", "Rp 96.000"]],
    },
    {
        "id": "pelanggan-order-kredit",
        "title": "Pelanggan, Order, Batas Kredit, dan Pemenuhan",
        "image": "04-pelanggan-order-kredit.png",
        "platform": "Web, Desktop Windows, dan Android Flutter",
        "role": "Sales, admin penjualan, kredit, gudang, dan pengiriman",
        "objective": "mengubah permintaan pelanggan menjadi order yang dapat dipenuhi, disetujui, disiapkan, dan dikirim dengan pengendalian kredit yang mudah dipahami",
        "regions": ["profil pelanggan", "alamat dan wilayah", "batas serta sisa kredit", "item order", "status persetujuan", "status siapkan dan kirim"],
        "steps": ["pilih pelanggan yang tepat", "konfirmasi alamat pengiriman", "periksa saldo piutang dan jatuh tempo", "masukkan item dan jumlah", "jalankan pemeriksaan harga dan kredit", "minta persetujuan bila diperlukan", "serahkan ke gudang", "pantau pengiriman sampai diterima"],
        "validations": ["pelanggan aktif dan tidak terblokir", "alamat pengiriman tersedia", "order tidak melampaui batas kredit tanpa persetujuan", "harga sesuai segmen", "stok dialokasikan satu kali", "status berubah berurutan", "penerima barang tercatat"],
        "risks": ["memilih pelanggan dengan nama mirip", "melewati pemeriksaan kredit", "mengubah order setelah gudang menyiapkan barang", "mengirim ke alamat lama", "membuat invoice sebelum barang siap", "menandai terkirim tanpa bukti"],
        "scenario": "Nofal membuat order untuk pelanggan lama yang memiliki sebagian piutang jatuh tempo. Sistem menahan order, manajemen menilai riwayat pembayaran, lalu menyetujui sebagian jumlah yang aman untuk dikirim.",
        "outcome": "Penjualan tetap berjalan tanpa menghilangkan disiplin kredit, sementara setiap keputusan pengecualian memiliki alasan dan jejak persetujuan.",
        "stats": [("Pelanggan aktif", "334"), ("Order diproses", "38"), ("Perlu persetujuan", "5"), ("Siap kirim", "17")],
        "table": ["Order", "Pelanggan", "Nilai", "Kredit", "Tahap"],
        "rows": [["SO-24081", "Toko Sehat", "Rp 2,8 jt", "Aman", "Siapkan"], ["SO-24082", "Apotek Melati", "Rp 4,1 jt", "Aman", "Kirim"], ["SO-24083", "Toko Maju", "Rp 6,2 jt", "Tinjau", "Persetujuan"], ["SO-24084", "Klinik Harmoni", "Rp 3,5 jt", "Aman", "Diterima"]],
    },
    {
        "id": "pembelian-penerimaan-hutang",
        "title": "Pembelian, Penerimaan Batch, dan Hutang Supplier",
        "image": "05-pembelian-penerimaan.png",
        "platform": "Web dan Desktop Windows",
        "role": "Pembelian, gudang, keuangan, dan manajemen",
        "objective": "merencanakan pembelian, menerima barang sesuai dokumen, mencatat batch dan kedaluwarsa, serta mengendalikan hutang supplier dari satu alur yang dapat diaudit",
        "regions": ["supplier dan syarat bayar", "nomor pesanan pembelian", "daftar barang", "hasil pemeriksaan penerimaan", "batch dan kedaluwarsa", "invoice serta jadwal hutang"],
        "steps": ["buat kebutuhan pembelian", "pilih supplier yang disetujui", "terbitkan pesanan pembelian", "cocokkan barang datang", "catat batch dan kedaluwarsa", "pisahkan barang bermasalah", "posting penerimaan", "verifikasi invoice", "jadwalkan pembayaran"],
        "validations": ["supplier aktif", "harga beli sesuai persetujuan", "jumlah diterima tidak melebihi toleransi", "batch tidak kosong", "tanggal kedaluwarsa memenuhi batas minimal", "invoice tidak ganda", "hutang sama dengan dokumen yang diterima"],
        "risks": ["mencatat barang yang belum diterima", "mencampur batch", "menerima kedaluwarsa terlalu dekat", "membayar invoice ganda", "mengubah harga setelah persetujuan", "menutup selisih tanpa berita acara"],
        "scenario": "Gudang menerima kiriman supplier dalam dua batch. Satu batch memenuhi syarat, sedangkan satu batch terlalu dekat dengan kedaluwarsa sehingga ditempatkan pada area karantina dan tidak menambah stok jual.",
        "outcome": "Stok hanya bertambah dari barang yang diterima dan lolos pemeriksaan, sementara hutang tercatat berdasarkan dokumen yang dapat dilacak kembali.",
        "stats": [("PO terbuka", "14"), ("Penerimaan hari ini", "8"), ("Karantina", "2 batch"), ("Hutang jatuh tempo", "Rp 32,7 jt")],
        "table": ["Supplier", "Dokumen", "Diterima", "Selisih", "Status"],
        "rows": [["PT APL", "PO-0712", "100%", "0", "Selesai"], ["Bina San Prima", "PO-0713", "92%", "8", "Sebagian"], ["Marga Nusantara", "PO-0714", "100%", "1 batch", "Karantina"], ["Duta Aman", "PO-0715", "-", "-", "Dalam perjalanan"]],
    },
    {
        "id": "stok-opname-expiry",
        "title": "Stok Opname, Penyesuaian, FEFO, dan Penarikan Produk",
        "image": "06-stok-opname.png",
        "platform": "Web, Desktop Windows, dan Android Flutter",
        "role": "Gudang, supervisor, quality control, dan auditor",
        "objective": "membandingkan stok fisik dengan sistem secara tertib, menjelaskan selisih, mengendalikan batch kedaluwarsa, dan menangani penarikan produk tanpa kehilangan jejak",
        "regions": ["lokasi dan sesi opname", "stok sistem", "hasil hitung fisik", "selisih dan nilai", "alasan penyesuaian", "status persetujuan serta posting"],
        "steps": ["tentukan ruang lingkup opname", "bekukan pergerakan bila diperlukan", "hitung tanpa melihat angka sistem", "masukkan hasil per batch", "ulang hitung selisih besar", "pilih alasan yang tepat", "minta persetujuan", "posting penyesuaian", "arsipkan berita acara"],
        "validations": ["lokasi dan batch benar", "petugas hitung dan pemeriksa berbeda", "selisih besar dihitung ulang", "alasan wajib terisi", "barang karantina tidak dianggap stok jual", "posting hanya sekali", "dokumen mempunyai waktu dan pengguna"],
        "risks": ["menghitung saat barang masih bergerak", "menggabungkan batch", "menyesuaikan tanpa pemeriksaan", "menghapus bukti selisih", "menjual batch yang ditarik", "mengubah tanggal kedaluwarsa agar lolos"],
        "scenario": "Tim gudang menemukan stok fisik Bodrex berbeda dua kotak dari sistem. Mereka menghitung ulang, memeriksa transaksi terakhir, menemukan kesalahan satuan, lalu mengajukan penyesuaian dengan bukti.",
        "outcome": "Saldo stok kembali dapat dipercaya, penyebab selisih diketahui, dan tindakan pencegahan dapat diterapkan pada proses penerimaan atau pengeluaran berikutnya.",
        "stats": [("Item dihitung", "524/626"), ("Selisih", "18"), ("Nilai selisih", "Rp 4,06 jt"), ("Perlu persetujuan", "6")],
        "table": ["Produk", "Batch", "Sistem", "Fisik", "Selisih"],
        "rows": [["Surya 12 RK", "SY-0726", "1", "0", "-1"], ["Oskadon", "OS-0626", "121", "119", "-2"], ["Paramex", "PR-0626", "707", "706", "-1"], ["Fatigon Spirit", "FS-0526", "0", "9", "+9"]],
    },
    {
        "id": "piutang-penagihan",
        "title": "Piutang, Pembayaran, dan Nota yang Dibawa Sales",
        "image": "07-piutang-penagihan.png",
        "platform": "Web, Desktop Windows, dan Android Flutter",
        "role": "Sales penagihan, kasir, keuangan, dan pemilik",
        "objective": "menentukan tagihan yang harus dibawa sales, mencatat pembayaran dengan bukti, mengalokasikan pembayaran ke invoice, dan memantau umur piutang secara akurat",
        "regions": ["daftar pelanggan", "invoice dan jatuh tempo", "umur piutang", "nota yang dibawa", "penerimaan pembayaran", "rekonsiliasi serta persetujuan"],
        "steps": ["pilih pelanggan atau wilayah", "urutkan invoice menurut jatuh tempo", "tentukan nota yang dibawa", "serahkan daftar kepada sales", "catat hasil penagihan", "unggah atau simpan bukti", "alokasikan pembayaran", "rekonsiliasi dengan kas atau bank", "tutup tugas penagihan"],
        "validations": ["invoice belum lunas", "jumlah dibawa sesuai saldo", "pembayaran tidak melebihi tagihan tanpa alasan", "metode pembayaran jelas", "bukti dapat dibaca", "tanggal penerimaan benar", "alokasi dan saldo akhir seimbang"],
        "risks": ["membawa nota tanpa serah terima", "mencatat pembayaran ke pelanggan salah", "menunda setoran", "menghapus pembayaran yang sudah direkonsiliasi", "menganggap transfer masuk tanpa verifikasi bank", "menutup invoice dengan potongan tanpa persetujuan"],
        "scenario": "Tohirin menerima daftar nota untuk wilayah C2, menagih tiga pelanggan, memperoleh dua pembayaran transfer dan satu janji bayar, lalu admin mencocokkan bukti dengan mutasi bank.",
        "outcome": "Pemilik mengetahui uang yang seharusnya diterima, yang benar-benar diterima, yang masih dibawa sales, dan alasan invoice belum selesai.",
        "stats": [("Total piutang", "Rp 99,6 jt"), ("Jatuh tempo", "Rp 18,9 jt"), ("Nota dibawa", "27"), ("Bayar hari ini", "Rp 12,4 jt")],
        "table": ["Pelanggan", "Invoice", "Jatuh tempo", "Saldo", "Tindakan"],
        "rows": [["Ani TK", "2606-0343", "11/07", "Rp 536 rb", "Verifikasi"], ["Maun TK", "2607-0216", "15/07", "Rp 1,03 jt", "Ditagih"], ["Lestari TK", "2606-0509", "15/07", "Rp 347 rb", "Dibayar"], ["Ema TK", "2606-0486", "14/07", "Rp 100 rb", "Janji bayar"]],
    },
    {
        "id": "laporan-tutup-periode",
        "title": "Laporan, Laba, Audit, Sinkronisasi, dan Tutup Periode",
        "image": "08-laporan-tutup-periode.png",
        "platform": "Web dan Desktop Windows",
        "role": "Pemilik, keuangan, auditor, dan administrator",
        "objective": "menghasilkan laporan yang dapat dijelaskan, memastikan seluruh perangkat tersinkron, melakukan rekonsiliasi, serta menutup periode tanpa menghapus sejarah transaksi",
        "regions": ["filter periode dan dimensi", "ringkasan laba kotor", "rekap per sales", "rekonsiliasi stok dan keuangan", "status sinkronisasi", "checklist persetujuan tutup periode"],
        "steps": ["tentukan periode laporan", "pastikan semua perangkat tersinkron", "periksa transaksi draft dan gagal", "rekonsiliasi stok", "rekonsiliasi piutang dan hutang", "tinjau laba per produk dan sales", "ekspor laporan", "buat backup", "minta persetujuan", "kunci periode"],
        "validations": ["tidak ada transaksi tertunda", "saldo subledger sama dengan laporan utama", "stok negatif telah dijelaskan", "retur masuk pada periode benar", "laba menggunakan HPP yang disetujui", "backup berhasil diuji", "penguncian tercatat dalam audit log"],
        "risks": ["menutup periode sebelum sinkronisasi", "menghapus transaksi lama", "mengubah laporan di luar sistem", "mengabaikan selisih kecil berulang", "membagikan file laporan tanpa perlindungan", "membuka kembali periode tanpa persetujuan"],
        "scenario": "Pada akhir bulan, admin memastikan semua order sales telah masuk, keuangan mencocokkan pembayaran, gudang menyelesaikan penyesuaian, lalu Muklis menyetujui penguncian setelah membaca laporan laba dan piutang.",
        "outcome": "Laporan bulan tersebut stabil, dapat dibandingkan dengan bulan berikutnya, dan setiap perubahan setelah penutupan memerlukan prosedur yang jelas serta dapat diaudit.",
        "stats": [("Omzet", "Rp 428,6 jt"), ("HPP", "Rp 367,2 jt"), ("Laba kotor", "Rp 61,4 jt"), ("Status periode", "Siap ditutup")],
        "table": ["Kontrol", "Penanggung jawab", "Hasil", "Bukti", "Status"],
        "rows": [["Sinkronisasi", "Admin", "0 tertunda", "Log", "Lulus"], ["Stok", "Gudang", "6 disetujui", "BA", "Lulus"], ["Piutang", "Keuangan", "Seimbang", "Rekap", "Lulus"], ["Backup", "Administrator", "Terverifikasi", "Checksum", "Lulus"]],
    },
]


def word_count(text: str) -> int:
    return len(re.findall(r"\b[\w-]+\b", text, flags=re.UNICODE))


def join_natural(values: list[str]) -> str:
    if len(values) == 1:
        return values[0]
    return ", ".join(values[:-1]) + ", dan " + values[-1]


def build_sections(chapter: dict) -> list[dict]:
    title = chapter["title"]
    objective = chapter["objective"]
    role = chapter["role"]
    platform = chapter["platform"]
    regions = chapter["regions"]
    steps = chapter["steps"]
    validations = chapter["validations"]
    risks = chapter["risks"]
    scenario = chapter["scenario"]
    outcome = chapter["outcome"]

    sections = [
        {
            "title": "Tujuan layar dan hasil yang diharapkan",
            "paragraphs": [
                f"Ilustrasi {title} menggambarkan ruang kerja yang dipakai untuk {objective}. Layar ini tidak dimaksudkan sebagai kumpulan angka yang hanya dilihat sekilas. Setiap bagian membantu pengguna menjawab pertanyaan operasional: apa yang sudah terjadi, apa yang belum selesai, siapa yang harus menindaklanjuti, dan bukti apa yang harus disimpan. Pengguna utama adalah {role}, sedangkan media yang digunakan adalah {platform}. Tampilan nyata dapat menyesuaikan ukuran perangkat dan hak akses, tetapi urutan berpikirnya tetap sama.",
                f"Hasil terbaik diperoleh ketika pengguna membuka layar dengan tujuan yang jelas. Sebelum menekan tombol apa pun, tentukan keputusan yang ingin dibuat dan periode data yang ingin diperiksa. Setelah pekerjaan selesai, hasil yang diharapkan adalah: {outcome} Dengan pendekatan ini, sistem bukan sekadar tempat mengetik, melainkan catatan bersama yang menghubungkan sales, gudang, administrasi, keuangan, dan pemilik. Setiap orang membaca sumber data yang sama sesuai batas kewenangannya.",
            ],
        },
        {
            "title": "Cara membaca ilustrasi dari atas ke bawah",
            "paragraphs": [
                f"Mulailah dari bagian paling atas. Di sana biasanya terdapat nama perusahaan, identitas pengguna, periode kerja, lokasi atau cabang, serta waktu pembaruan terakhir. Informasi ini wajib dibaca agar pengguna tidak mengolah data milik tenant, periode, atau gudang yang salah. Setelah identitas benar, lanjutkan ke bagian utama yang terdiri atas {join_natural(regions)}. Susunan tersebut dibuat agar pengguna melihat ringkasan lebih dulu, kemudian beralih ke rincian dan tindakan.",
                f"Warna pada ilustrasi berfungsi sebagai petunjuk, bukan hiasan. Hijau menunjukkan kondisi selesai atau aman, kuning menunjukkan keadaan yang perlu diperiksa, merah menunjukkan risiko atau pekerjaan yang tertahan, dan abu-abu menunjukkan informasi pendukung. Jangan mengambil keputusan hanya berdasarkan warna; buka rincian yang mendasarinya. Angka pada ilustrasi adalah contoh pembelajaran dan bukan data produksi Caruban Medika Nusantara. Pada aplikasi asli, pengguna harus memeriksa nomor dokumen, tanggal, pihak terkait, dan status sebelum melanjutkan.",
            ],
        },
        {
            "title": "Persiapan sebelum memulai pekerjaan",
            "paragraphs": [
                f"Sebelum menggunakan {title}, pastikan akun yang dipakai adalah akun pribadi dan peran yang tampil sesuai pekerjaan. Jangan berbagi akun karena sistem mencatat siapa yang membuat, mengubah, menyetujui, atau membatalkan data. Periksa koneksi, indikator sinkronisasi, tanggal serta jam perangkat, dan cabang aktif. Untuk pekerjaan yang memengaruhi stok atau uang, siapkan dokumen pendukung seperti pesanan, invoice, bukti penerimaan, bukti transfer, atau berita acara. Persiapan sederhana ini mencegah kesalahan yang sulit dilacak kemudian.",
                f"Pengguna juga perlu memahami batas kewenangan. {role} tidak selalu mempunyai hak yang sama. Bila tombol tidak terlihat atau tidak aktif, jangan mencari cara untuk melewati pengendalian. Catat nomor dokumen dan minta bantuan atasan atau admin hak akses. Data sensitif tidak boleh dikirim melalui grup umum. Ketika bekerja menggunakan {platform}, pastikan layar terkunci saat ditinggalkan dan berkas hasil ekspor disimpan pada tempat yang disetujui perusahaan.",
            ],
        },
        {
            "title": "Langkah kerja utama",
            "paragraphs": [
                f"Urutan kerja yang disarankan adalah {join_natural(steps)}. Jalankan urutan tersebut secara tertib karena setiap langkah menyiapkan informasi bagi langkah berikutnya. Misalnya, pemilihan pelanggan atau supplier menentukan alamat, syarat bayar, harga, dan batas kredit; pemilihan produk menentukan satuan, batch, HPP, dan ketersediaan. Jika data awal salah, hasil akhir juga salah meskipun perhitungan sistem benar. Gunakan pencarian dan filter untuk mempersempit data, lalu buka satu catatan yang benar sebelum melakukan tindakan.",
                f"Sesudah memasukkan data, baca kembali ringkasan sebelum menyimpan. Sistem dapat memberikan peringatan, menahan proses, atau meminta persetujuan. Peringatan bukan gangguan; peringatan menjelaskan risiko yang harus diselesaikan. Jika jaringan terputus, periksa apakah pekerjaan tersimpan sebagai draft lokal, antrean sinkronisasi, atau belum tersimpan sama sekali. Jangan mengulang transaksi tanpa memeriksa nomor referensi karena pengulangan dapat menghasilkan data ganda. Setelah berhasil, catat status akhir dan pihak yang menerima pekerjaan berikutnya.",
            ],
        },
        {
            "title": "Pemeriksaan sebelum menyimpan atau menyetujui",
            "paragraphs": [
                f"Gunakan daftar pemeriksaan berikut sebagai kebiasaan: {join_natural(validations)}. Pemeriksaan tidak perlu memakan waktu lama jika dilakukan pada setiap transaksi. Fokus pada identitas pihak, tanggal, mata uang, satuan, jumlah, harga, pajak, potongan, batch, jatuh tempo, dan status. Untuk transaksi bernilai besar atau berisiko, lakukan prinsip empat mata: satu orang menyiapkan dan orang lain memeriksa. Pemisahan ini melindungi pegawai sekaligus perusahaan dari kesalahan yang tidak disengaja.",
                f"Setelah tombol simpan ditekan, cari tanda keberhasilan yang nyata, misalnya nomor dokumen, status tersimpan, waktu sinkronisasi, atau perubahan pada daftar. Pesan singkat saja tidak cukup jika layar kemudian menunjukkan status gagal. Pada perangkat Android, bedakan antara tersimpan di perangkat dan sudah diterima server. Pada Web atau Desktop, jangan menutup jendela saat proses masih berjalan. Bila hasil tidak sesuai, jangan membuat koreksi diam-diam; gunakan prosedur edit, pembatalan, retur, atau penyesuaian yang menyediakan alasan dan jejak audit.",
            ],
        },
        {
            "title": "Kesalahan umum dan cara mencegahnya",
            "paragraphs": [
                f"Kesalahan yang paling sering terjadi pada alur ini adalah {join_natural(risks)}. Sebagian besar bukan disebabkan oleh kerusakan sistem, melainkan karena konteks tidak diperiksa, pekerjaan terburu-buru, atau tanggung jawab tidak jelas. Cara pencegahannya adalah menggunakan akun sendiri, memulai dari dokumen sumber, membaca ringkasan, dan menyelesaikan satu transaksi sebelum pindah ke transaksi lain. Hindari membuka banyak tab untuk pekerjaan yang sama karena pengguna dapat menyimpan versi yang berbeda.",
                f"Jika kesalahan sudah terjadi, hentikan proses berikutnya dan catat fakta apa adanya: nomor dokumen, waktu, pengguna, nilai, status, dan dampak. Jangan menghapus bukti atau mengganti data agar tampak benar. Laporkan kepada penanggung jawab, pilih mekanisme koreksi yang tersedia, lalu periksa laporan setelah koreksi diposting. Sikap terbuka terhadap kesalahan mempercepat pemulihan dan membuat aturan kerja dapat diperbaiki. Sistem menyimpan jejak perubahan untuk membantu penelusuran, bukan untuk menyalahkan pengguna yang mengikuti prosedur.",
            ],
        },
        {
            "title": "Pembagian tanggung jawab antarperan",
            "paragraphs": [
                f"Pada {title}, pembuat data bertanggung jawab terhadap kelengkapan informasi awal; pemeriksa bertanggung jawab terhadap kesesuaian dokumen; penyetuju bertanggung jawab terhadap keputusan bisnis; dan administrator bertanggung jawab terhadap ketersediaan sistem serta hak akses. Pemilik usaha menggunakan hasil untuk mengarahkan bisnis, tetapi tidak seharusnya mengambil alih seluruh pekerjaan operasional. Pembagian tugas yang jelas mempercepat proses karena setiap orang mengetahui kapan harus bertindak dan kapan harus menyerahkan pekerjaan.",
                f"Gunakan kolom catatan, status, dan penanggung jawab untuk melakukan serah terima. Instruksi lisan tetap boleh digunakan untuk mempercepat komunikasi, tetapi keputusan penting harus dicatat di sistem atau dokumen resmi. Jika seorang pegawai tidak hadir, atasan menunjuk pengganti melalui hak akses sementara yang terdokumentasi. Jangan memberikan kata sandi pegawai yang tidak hadir. Setelah tugas selesai, cabut hak sementara dan tinjau aktivitas yang dilakukan selama masa penggantian.",
            ],
        },
        {
            "title": "Perbedaan penggunaan Web, Windows, dan Android",
            "paragraphs": [
                f"{platform} menyediakan kemampuan yang sama pada inti data, tetapi cara berinteraksinya dapat berbeda. Web dan Desktop cocok untuk tabel lebar, pemeriksaan banyak dokumen, rekonsiliasi, ekspor, serta persetujuan. Android cocok untuk pekerjaan lapangan, pencarian cepat, pemindaian, pengambilan bukti, dan pekerjaan yang perlu tetap berjalan ketika koneksi tidak stabil. Pengguna tidak perlu memaksakan semua pekerjaan ke satu perangkat. Pilih perangkat berdasarkan risiko, jumlah data, dan kondisi tempat kerja.",
                f"Pada layar kecil, bagian informasi dapat disusun vertikal dan beberapa kolom diringkas menjadi kartu. Gulir sampai akhir dan buka detail sebelum menyimpulkan bahwa data tidak tersedia. Pada layar besar, jangan mengabaikan kolom di sebelah kanan; gunakan pengguliran horizontal bila diperlukan. Perubahan yang dibuat di satu perangkat baru terlihat di perangkat lain setelah sinkronisasi berhasil. Karena itu, indikator pembaruan terakhir harus menjadi bagian dari pemeriksaan, terutama sebelum menyetujui stok, kredit, pembayaran, atau penutupan periode.",
            ],
        },
        {
            "title": "Contoh kasus sederhana",
            "paragraphs": [
                f"{scenario} Dalam contoh ini, pengguna tidak langsung mengejar kecepatan. Ia memeriksa identitas, konteks, dokumen, dan status. Ketika menemukan peringatan, ia membuka rincian dan menentukan apakah pekerjaan dapat dilanjutkan sendiri atau membutuhkan persetujuan. Seluruh tindakan diberi nomor referensi sehingga pihak berikutnya tidak perlu menebak asal informasi. Cara kerja seperti ini membuat pelayanan pelanggan tetap cepat tanpa mengorbankan pengendalian.",
                f"Setelah kasus selesai, pengguna membandingkan hasil dengan tujuan awal. Ia memastikan tidak ada draft tertinggal, antrean sinkronisasi kosong atau diketahui, dan laporan terkait telah berubah sesuai transaksi. Bila hasil belum muncul, ia tidak langsung mengulangi input; ia memeriksa status dan log lebih dahulu. Contoh ini dapat dipakai sebagai latihan. Gunakan data demo, lakukan langkah dari awal sampai akhir, lalu minta peserta lain menjelaskan apa yang terjadi dengan bahasa mereka sendiri.",
            ],
        },
        {
            "title": "Bukti, audit, dan keamanan informasi",
            "paragraphs": [
                f"Setiap pekerjaan pada {title} menghasilkan bukti. Bukti dapat berupa nomor dokumen, waktu pembuatan, pengguna, status, foto, tanda terima, file ekspor, atau catatan persetujuan. Bukti harus cukup untuk menjawab siapa, apa, kapan, di mana, mengapa, dan berapa nilainya. Simpan bukti pada fasilitas yang disediakan sistem atau lokasi resmi perusahaan. Jangan menggunakan foto pribadi, penyimpanan umum, atau aplikasi pesan sebagai arsip utama.",
                f"Hak akses mengikuti kebutuhan kerja minimum. Data harga beli, laba, saldo piutang, nomor rekening, dan identitas pelanggan hanya boleh dibuka oleh peran yang memerlukannya. Ketika menampilkan layar kepada pihak luar, gunakan mode demo atau samarkan data. Saat mengekspor laporan, periksa penerima dan masa simpan. Jika perangkat hilang, segera laporkan agar sesi dapat dicabut. Audit log tidak boleh dihapus untuk menyembunyikan kesalahan; justru catatan tersebut membantu perusahaan memahami dan memperbaiki proses.",
            ],
        },
        {
            "title": "Checklist penyelesaian pekerjaan",
            "paragraphs": [
                f"Sebelum meninggalkan layar, tanyakan: apakah tujuan sudah tercapai, apakah nomor dokumen sudah terbentuk, apakah status akhir benar, apakah data sudah tersinkron, apakah dokumen pendukung tersimpan, apakah pihak berikutnya sudah menerima tugas, dan apakah terdapat peringatan yang belum ditangani. Jika satu jawaban masih tidak jelas, pekerjaan belum benar-benar selesai. Tuliskan catatan singkat agar pergantian shift atau pemeriksaan berikutnya tidak dimulai dari nol.",
                f"Pada akhir hari, supervisor meninjau daftar transaksi gagal, draft, pembatalan, penyesuaian, dan persetujuan tertunda. Tinjauan harian lebih ringan daripada menunggu akhir bulan. Gunakan laporan pengecualian untuk mencari hal yang tidak biasa, bukan hanya laporan total. Apabila seluruh pemeriksaan telah lulus, pengguna dapat menutup tugas dan melanjutkan pekerjaan lain. Kebiasaan menyelesaikan satu siklus penuh adalah dasar data yang akurat dan laporan pemilik yang dapat dipercaya.",
            ],
        },
        {
            "title": "Cara melatih pengguna baru",
            "paragraphs": [
                f"Pelatihan dimulai dengan menjelaskan tujuan bisnis, bukan nama tombol. Tunjukkan ilustrasi, sebutkan peran {role}, lalu minta peserta menunjuk {join_natural(regions)}. Setelah itu demonstrasikan satu kasus normal dan satu kasus yang menghasilkan peringatan. Peserta kemudian mengulang menggunakan data demo sambil menjelaskan alasan setiap langkah. Pendamping hanya membantu ketika peserta berhenti, sehingga pemahaman tidak bergantung pada hafalan.",
                f"Pengguna dinyatakan siap ketika mampu menyelesaikan kasus, menemukan kesalahan yang sengaja disisipkan, menjelaskan bukti yang harus disimpan, dan mengetahui kapan harus meminta persetujuan. Catat hasil latihan dan topik yang perlu diulang. Jangan melatih langsung dengan data produksi yang sensitif. Setelah pengguna mulai bekerja, lakukan pendampingan pada transaksi pertama dan tinjau hasilnya. Pelatihan yang baik membuat orang nonteknis percaya diri karena mereka memahami maksud proses, bukan sekadar mengikuti layar.",
            ],
        },
    ]
    total = word_count(" ".join(p for section in sections for p in section["paragraphs"]))
    if total < 1500:
        raise ValueError(f"{chapter['id']} hanya memiliki {total} kata; minimum 1500")
    return sections


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    name = "segoeuib.ttf" if bold else "segoeui.ttf"
    return ImageFont.truetype(str(Path("C:/Windows/Fonts") / name), size=size)


def rounded(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], fill: str, outline: str = "#D8E0EA", radius: int = 14, width: int = 2) -> None:
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def fit_text(draw: ImageDraw.ImageDraw, text: str, box: tuple[int, int, int, int], size: int, color: str, bold: bool = False, line_gap: int = 6) -> None:
    x1, y1, x2, y2 = box
    current = font(size, bold)
    max_width = x2 - x1
    words = text.split()
    lines: list[str] = []
    line = ""
    for word in words:
        trial = f"{line} {word}".strip()
        if draw.textbbox((0, 0), trial, font=current)[2] <= max_width:
            line = trial
        else:
            if line:
                lines.append(line)
            line = word
    if line:
        lines.append(line)
    y = y1
    for value in lines:
        draw.text((x1, y), value, font=current, fill=color)
        y += size + line_gap
        if y > y2:
            break


def generate_illustration(chapter: dict) -> None:
    canvas = Image.new("RGB", (1600, 900), "#F4F7FB")
    draw = ImageDraw.Draw(canvas)
    draw.rectangle((0, 0, 1600, 82), fill="#0F172A")
    draw.rounded_rectangle((30, 20, 72, 62), radius=10, fill="#0F766E")
    draw.text((43, 27), "eB", font=font(17, True), fill="white")
    draw.text((90, 25), "Caruban Medika Nusantara", font=font(22, True), fill="white")
    draw.text((90, 52), chapter["platform"], font=font(12), fill="#CBD5E1")
    draw.rounded_rectangle((1235, 20, 1565, 62), radius=12, fill="#1E293B", outline="#334155")
    draw.text((1260, 31), "Data terakhir: hari ini 08.42", font=font(14), fill="#D1FAE5")

    draw.rectangle((0, 82, 220, 900), fill="#FFFFFF")
    draw.text((28, 112), "NAVIGASI", font=font(12, True), fill="#64748B")
    navs = ["Beranda", "Penjualan", "Produk & stok", "Pelanggan", "Pembelian", "Keuangan", "Laporan"]
    selected = min(6, CHAPTERS.index(chapter))
    for index, nav in enumerate(navs):
        top = 150 + index * 58
        if index == selected:
            draw.rounded_rectangle((18, top - 8, 202, top + 34), radius=9, fill="#E6FFFA")
            draw.rectangle((18, top - 8, 24, top + 34), fill="#0F766E")
        draw.text((38, top), nav, font=font(15, index == selected), fill="#0F172A")
    draw.text((28, 828), "Bantuan", font=font(14), fill="#475569")
    draw.text((28, 858), "Keluar", font=font(14), fill="#475569")

    fit_text(draw, chapter["title"], (260, 112, 1110, 180), 31, "#0F172A", True)
    draw.rounded_rectangle((1250, 111, 1560, 153), radius=12, fill="#ECFDF5", outline="#A7F3D0")
    draw.text((1273, 123), "ILUSTRASI KONSEPTUAL", font=font(14, True), fill="#047857")
    draw.text((260, 184), "Contoh tampilan untuk pelatihan. Angka di bawah bukan data produksi.", font=font(14), fill="#64748B")

    stats = chapter["stats"]
    for index, (label, value) in enumerate(stats):
        x = 260 + index * 320
        rounded(draw, (x, 225, x + 295, 330), "#FFFFFF")
        draw.text((x + 20, 245), label.upper(), font=font(12, True), fill="#64748B")
        draw.text((x + 20, 278), value, font=font(25, True), fill="#0F172A")

    rounded(draw, (260, 360, 1245, 820), "#FFFFFF")
    draw.text((285, 382), "Rincian pekerjaan", font=font(18, True), fill="#0F172A")
    headers = chapter["table"]
    rows = chapter["rows"]
    col_widths = [220, 190, 170, 170, 165]
    x_positions = [285]
    for width in col_widths[:-1]:
        x_positions.append(x_positions[-1] + width)
    draw.rounded_rectangle((280, 425, 1220, 470), radius=8, fill="#E8EEF5")
    for index, value in enumerate(headers):
        draw.text((x_positions[index], 438), value, font=font(13, True), fill="#334155")
    for row_index, row in enumerate(rows):
        y = 485 + row_index * 68
        if row_index % 2:
            draw.rounded_rectangle((280, y - 10, 1220, y + 42), radius=5, fill="#F8FAFC")
        for index, value in enumerate(row):
            draw.text((x_positions[index], y), value, font=font(14, index == 0), fill="#0F172A")
        draw.line((280, y + 48, 1220, y + 48), fill="#E2E8F0", width=1)

    rounded(draw, (1270, 360, 1560, 820), "#0F172A", outline="#0F172A")
    draw.text((1295, 386), "LANGKAH BERIKUTNYA", font=font(13, True), fill="#5EEAD4")
    for index, step in enumerate(chapter["steps"][:5], start=1):
        y = 430 + (index - 1) * 73
        draw.ellipse((1294, y, 1324, y + 30), fill="#0F766E")
        draw.text((1304, y + 5), str(index), font=font(13, True), fill="white")
        fit_text(draw, step.capitalize(), (1338, y - 2, 1538, y + 57), 13, "#E2E8F0", False, 4)
    draw.rounded_rectangle((1293, 766, 1537, 804), radius=9, fill="#14B8A6")
    draw.text((1354, 776), "Buka rincian", font=font(14, True), fill="#042F2E")
    canvas.save(IMAGES / chapter["image"], quality=96)


def set_repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    element = OxmlElement("w:tblHeader")
    element.set(qn("w:val"), "true")
    properties.append(element)


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    element = OxmlElement("w:shd")
    element.set(qn("w:fill"), fill)
    properties.append(element)


def set_table_widths(table, widths: list[int]) -> None:
    table.autofit = False
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    properties = table._tbl.tblPr
    width_element = properties.first_child_found_in("w:tblW")
    if width_element is None:
        width_element = OxmlElement("w:tblW")
        properties.insert(0, width_element)
    width_element.set(qn("w:w"), str(sum(widths)))
    width_element.set(qn("w:type"), "dxa")
    indent = properties.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell_properties = cell._tc.get_or_add_tcPr()
            cell_width = cell_properties.first_child_found_in("w:tcW")
            cell_width.set(qn("w:w"), str(widths[index]))
            cell_width.set(qn("w:type"), "dxa")


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, end])


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Title", 30, INK, 0, 8),
        ("Heading 1", 16, RGBColor(46, 116, 181), 18, 10),
        ("Heading 2", 13, RGBColor(46, 116, 181), 14, 7),
        ("Heading 3", 12, RGBColor(31, 77, 120), 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "PANDUAN OPERASIONAL BERGAMBAR  |  eBisnis Inventory / Sales"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(8)
        run.font.color.rgb = MUTED
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Caruban Medika Nusantara  |  Halaman ")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED
    add_page_number(footer)


def add_cover(document: Document, data: dict) -> None:
    document.add_paragraph().paragraph_format.space_after = Pt(54)
    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("WEB  |  WINDOWS  |  ANDROID FLUTTER")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = TEAL
    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(data["meta"]["title"])
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(22)
    run = subtitle.add_run(data["meta"]["subtitle"])
    run.font.size = Pt(14)
    run.font.color.rgb = MUTED
    picture = document.add_paragraph()
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = picture.add_run()
    inline = run.add_picture(str(IMAGES / data["chapters"][0]["image"]), width=Inches(6.45))
    inline._inline.docPr.set("descr", "Ilustrasi konseptual dashboard pemilik Caruban Medika Nusantara")
    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(12)
    note.add_run("Bahasa formal yang mudah dipahami pengguna nonteknis").bold = True
    meta = document.add_table(rows=2, cols=3)
    meta.style = "Table Grid"
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    labels = [("VERSI", data["meta"]["version"]), ("DIPERBARUI", data["meta"]["updated"]), ("BAB BERGAMBAR", str(len(data["chapters"])))]
    for index, (label, value) in enumerate(labels):
        meta.cell(0, index).text = label
        meta.cell(1, index).text = value
        set_cell_shading(meta.cell(0, index), "E8EEF5")
        for run in meta.cell(0, index).paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = TEAL
    set_repeat_table_header(meta.rows[0])
    set_table_widths(meta, [3120, 3120, 3120])
    document.add_page_break()


def add_front_matter(document: Document, data: dict) -> None:
    document.add_heading("Cara Menggunakan Panduan Ini", level=1)
    document.add_paragraph(
        "Panduan ini adalah volume operasional bergambar yang melengkapi panduan ringkas Inventory / Sales. Setiap bab dimulai dengan ilustrasi layar, kemudian dijelaskan secara formal dan bertahap untuk pengguna nonteknis. Angka, nama pelanggan, dan status pada ilustrasi hanyalah contoh pelatihan. Jangan menggunakannya sebagai dasar transaksi produksi."
    )
    document.add_paragraph(
        "Baca bab yang sesuai dengan pekerjaan Anda. Ikuti urutan dari tujuan layar, persiapan, langkah kerja, pemeriksaan, penanganan kesalahan, pembagian peran, contoh kasus, audit, checklist, hingga cara melatih pengguna baru. Setiap bab mempunyai sedikitnya 1.500 kata agar alasan di balik proses tidak hilang."
    )
    document.add_heading("Daftar Bab Bergambar", level=1)
    for index, chapter in enumerate(data["chapters"], start=1):
        paragraph = document.add_paragraph(style="List Number")
        paragraph.add_run(f"{chapter['title']} ({chapter['wordCount']:,} kata)")
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["Tanda", "Arti", "Tindakan pengguna"]
    for index, label in enumerate(headers):
        table.cell(0, index).text = label
        set_cell_shading(table.cell(0, index), "E8EEF5")
        table.cell(0, index).paragraphs[0].runs[0].bold = True
    for values in [
        ("Hijau", "Selesai atau aman", "Tetap buka rincian untuk memastikan bukti."),
        ("Kuning", "Perlu diperiksa", "Baca peringatan dan selesaikan penyebabnya."),
        ("Merah", "Tertahan atau berisiko", "Jangan lanjut sebelum ada penyelesaian atau persetujuan."),
        ("Abu-abu", "Informasi pendukung", "Gunakan sebagai konteks, bukan keputusan tunggal."),
    ]:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
    set_repeat_table_header(table.rows[0])
    set_table_widths(table, [1500, 2600, 5260])
    document.add_page_break()


def add_chapter(document: Document, index: int, chapter: dict) -> None:
    if index > 1:
        document.add_page_break()
    document.add_heading(f"{index}. {chapter['title']}", level=1)
    lead = document.add_paragraph()
    lead_run = lead.add_run(chapter["summary"])
    lead_run.bold = True
    lead_run.font.color.rgb = MUTED
    picture = document.add_paragraph()
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    inline = picture.add_run().add_picture(str(IMAGES / chapter["image"]), width=Inches(6.45))
    inline._inline.docPr.set("descr", chapter["alt"])
    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(10)
    run = caption.add_run(f"Gambar {index}. {chapter['alt']} Ilustrasi konseptual, bukan data produksi.")
    run.italic = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = MUTED
    callout = document.add_table(rows=1, cols=1)
    callout.style = "Table Grid"
    set_cell_shading(callout.cell(0, 0), "ECFDF5")
    callout.cell(0, 0).text = f"Cakupan penjelasan gambar: {chapter['wordCount']:,} kata. Platform: {chapter['platform']}. Pengguna utama: {chapter['role']}."
    set_repeat_table_header(callout.rows[0])
    set_table_widths(callout, [PAGE_WIDTH_DXA])
    for section in chapter["sections"]:
        document.add_heading(section["title"], level=2)
        for text in section["paragraphs"]:
            document.add_paragraph(text)
    document.add_heading("Ringkasan langkah kerja", level=2)
    for step in chapter["steps"]:
        document.add_paragraph(step.capitalize(), style="List Number")
    document.add_heading("Pemeriksaan penerimaan pengguna", level=2)
    for validation in chapter["validations"]:
        document.add_paragraph(validation.capitalize(), style="List Bullet")


def build() -> dict:
    PUBLIC.mkdir(parents=True, exist_ok=True)
    IMAGES.mkdir(parents=True, exist_ok=True)
    REPORT.parent.mkdir(parents=True, exist_ok=True)
    chapters = []
    counts = {}
    for source in CHAPTERS:
        generate_illustration(source)
        sections = build_sections(source)
        count = word_count(" ".join(p for section in sections for p in section["paragraphs"]))
        counts[source["id"]] = count
        chapters.append(
            {
                "id": source["id"],
                "title": source["title"],
                "image": source["image"],
                "alt": f"Ilustrasi layar {source['title']} pada eBisnis Inventory / Sales.",
                "platform": source["platform"],
                "role": source["role"],
                "summary": source["objective"].capitalize() + ".",
                "steps": source["steps"],
                "validations": source["validations"],
                "sections": sections,
                "wordCount": count,
            }
        )
    data = {
        "meta": {
            "title": "Panduan Operasional Bergambar eBisnis Inventory / Sales",
            "subtitle": "Penjelasan end to end untuk pemilik, admin, sales, gudang, keuangan, dan pengguna nonteknis",
            "version": "2.0",
            "updated": "5 Agustus 2026",
        },
        "chapters": chapters,
    }
    CONTENT.write_text(json.dumps(data, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    REPORT.write_text(json.dumps({"minimumWordsPerIllustration": 1500, "counts": counts}, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")

    document = Document()
    configure_document(document)
    add_cover(document, data)
    add_front_matter(document, data)
    for index, chapter in enumerate(chapters, start=1):
        add_chapter(document, index, chapter)
    properties = document.core_properties
    properties.title = data["meta"]["title"]
    properties.subject = "Panduan operasional bergambar Web, Windows, dan Android Flutter"
    properties.author = "eBisnis.id"
    properties.keywords = "inventory, sales, user manual, flutter, android, windows, web, Caruban Medika Nusantara"
    document.save(DOCX)
    return {"docx": str(DOCX), "content": str(CONTENT), "counts": counts}


if __name__ == "__main__":
    print(json.dumps(build(), ensure_ascii=False, indent=2))
